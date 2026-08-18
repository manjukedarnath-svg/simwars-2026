from flask import Flask, request, send_file, render_template_string, jsonify, redirect, session
import io
import os
import re
import sqlite3
from urllib.parse import quote
from datetime import datetime, timedelta
from docx import Document

DB_PATH = os.environ.get('DB_PATH', os.path.join(os.path.dirname(__file__), 'scores.db'))

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    with get_db() as conn:
        conn.execute('''CREATE TABLE IF NOT EXISTS scores
                        (key TEXT PRIMARY KEY, value TEXT, updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')
        conn.commit()

init_db()

def _get_weasyprint():
    try:
        from weasyprint import HTML as WeasyHTML
        return WeasyHTML
    except Exception:
        return None

app = Flask(__name__)

import secrets as _secrets
# Secret key: use env var if set; otherwise random per boot. Random fallback means all
# sessions log out on redeploy — safe default now that the repo is public. Set
# FLASK_SECRET_KEY in Railway → Variables to keep sessions across deploys.
app.secret_key = os.environ.get('FLASK_SECRET_KEY') or _secrets.token_hex(32)

ORGANISER_PASSWORD = '@MSPtrio2023sim'
JUDGE_PASSWORD = '@SIMblore2014'
TEAM_PASSWORD_PREFIX = '@SIMwarsTEAM'  # full password = prefix + draw number, e.g. @SIMwarsTEAM7 — also used by that team's mentor/director

def _ist_now():
    return datetime.utcnow() + timedelta(hours=5, minutes=30)

# Scoreboard opens per role (IST): organisers immediately, judges 21 Aug, teams event day 22 Aug
SCOREBOARD_OPENS = {
    'organiser': None,
    'judge': datetime(2026, 8, 21, 0, 0),
    'team': datetime(2026, 8, 22, 0, 0),
}

UNLOCK_TEMPLATE = '''
<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>SimWars — Enter Password</title>
<style>
body{font-family:-apple-system,Segoe UI,Roboto,sans-serif;background:#0f172a;color:#fff;display:flex;align-items:center;justify-content:center;min-height:100vh;margin:0;}
.card{background:#1e293b;padding:32px 36px;border-radius:14px;max-width:360px;width:90%;box-shadow:0 10px 40px rgba(0,0,0,.4);}
h1{font-size:1.1rem;margin:0 0 6px;}
p{color:#94a3b8;font-size:0.85rem;margin:0 0 18px;}
input{width:100%;box-sizing:border-box;padding:10px 12px;border-radius:8px;border:1px solid #334155;background:#0f172a;color:#fff;font-size:0.95rem;margin-bottom:12px;}
button{width:100%;padding:10px;border:none;border-radius:8px;background:#7C3AED;color:#fff;font-weight:700;cursor:pointer;font-size:0.9rem;}
.err{color:#f87171;font-size:0.82rem;margin-bottom:10px;}
</style></head>
<body>
<form class="card" method="post">
<h1>🔒 Organisers & Judges Only</h1>
<p>Enter your Organiser or Judge password to continue.</p>
{% if error %}<div class="err">Incorrect password. Try again.</div>{% endif %}
<input type="password" name="password" placeholder="Password" autofocus autocomplete="off">
<button type="submit">Unlock</button>
</form>
</body></html>
'''

def require_role():
    if session.get('role') not in ('organiser', 'judge'):
        full_path = request.path
        if request.query_string:
            full_path += '?' + request.query_string.decode()
        return redirect('/unlock?next=' + quote(full_path, safe=''))
    return None

def require_organiser():
    """Stricter than require_role(): only the organiser password grants access.
    Used for the full case scenario library (/library, /case/<id>) — judges get
    a scenario summary on their own scoring sheet instead, not the full script."""
    if session.get('role') != 'organiser':
        full_path = request.path
        if request.query_string:
            full_path += '?' + request.query_string.decode()
        return redirect('/unlock?next=' + quote(full_path, safe=''))
    return None

@app.route('/unlock', methods=['GET', 'POST'])
def unlock():
    next_url = request.args.get('next') or request.form.get('next') or '/library'
    if request.method == 'POST':
        pw = request.form.get('password', '').strip()
        if pw == ORGANISER_PASSWORD:
            session['role'] = 'organiser'
            return redirect(next_url)
        elif pw == JUDGE_PASSWORD:
            session['role'] = 'judge'
            return redirect(next_url)
        return render_template_string(UNLOCK_TEMPLATE, error=True)
    return render_template_string(UNLOCK_TEMPLATE, error=False)

# ============================================
# CASE DATA (All 8 Cases)
# ============================================
CASES = [
    {
        "id": "B1",
        "round": "BLS CASES",
        "title": "Near Drowning — Submersion Arrest with Severe Hypothermia",
        "summary": "12-year-old female (30 kg) pulled from a lake after ~10 min submersion, ~15 min total CPR, core temp 33°C, fine VF on the monitor. EMS have already given TWO rounds of Adrenaline but no shock. The team must defibrillate FIRST, refuse to stop CPR for fixed dilated pupils ('not dead until warm and dead'), actively rewarm, and state ECMO referral criteria for refractory VF.",
        "background": "Vaibhavi, 12 years old, 30 kg, pulled out of the lake about 15 minutes ago — under water around 10 minutes. Bystander CPR within a couple of minutes, continued by EMS. Not breathing, no pulse. Two rounds of Adrenaline given. Temperature 33°C. Pupils fixed and dilated bilaterally. IV in.",
        "expanded_history": "Family day trip; went into open lake water to retrieve a ball, submerged before anyone noticed — ~10 min based on last seen vs found. Bystander CPR within ~2 min of retrieval. Father (who pulled her out) is at the bedside, distressed; mother not present. Previously fit and well, no allergies, no seizure or cardiac history. EMS written note: HR 0, T 33°C tympanic on scene.",
        "stages": [
            {"name": "Stage 1: Initial Assessment (0:00–2:00)", "vitals": "HR: unrecordable (fine VF), RR: 0, BP: unrecordable, SpO2: unrecordable, T: 33°C, GCS 3, pupils fixed & dilated bilaterally", "condition": "Pulseless, apnoeic, in arrest on arrival. Bystander then EMS CPR in progress, total downtime ~15 min. Two rounds of Adrenaline already given, NO defibrillation yet.", "expected": "Recognise arrest; continue high-quality CPR without interrupting for pupil assessment. Call for help, allocate roles. Rhythm check → VF → DEFIBRILLATE 2 J/kg immediately (do not wait for or repeat Adrenaline). Request core temperature. Expose, dry, begin active rewarming — warm IV fluids, warm humidified O2, forced-air warming.", "notes": "Shock within the first 60 s → Stage 2 with VF persisting (expected). Further Adrenaline instead of shocking, or CPR paused to declare death on fixed pupils → remain in Stage 1, nurse escalates at 3:00."},
            {"name": "Stage 2: CPR Cycles / Rewarming (2:00–5:00)", "vitals": "Rhythm: VF persists on re-check. T static at 33°C without active rewarming; rising ~0.2–0.3°C/min with it. K 4.5 (on request)", "condition": "VF refractory to initial shock(s) — the expected consequence of the hypothermic myocardium, not a team failure.", "expected": "2-minute CPR cycles, defibrillating each cycle, escalating to 4 J/kg. Amiodarone 5 mg/kg IV after the 3rd shock. Multi-modality active rewarming continued/escalated. States ECMO referral criteria once K available — witnessed arrest, hypothermia, K <12.", "notes": "Active rewarming + good CPR + correct shock/drug sequence → ROSC ~5:30 (Stage 3). Inadequate rewarming + poor CPR → routes toward Stage 5 (Rescue)."},
            {"name": "Stage 3: ROSC (≈5:30)", "vitals": "HR 88 sinus, BP 62/40, T 33°C rising, GCS 4", "condition": "Return of spontaneous circulation — scripted once active rewarming, correct shock sequence and good CPR are delivered.", "expected": "Recognise ROSC promptly, stop CPR, confirm pulse. Secure airway (Ketamine + Rocuronium). Continue rewarming toward 35°C. Post-ROSC 12-lead. Avoid hyperoxia and hyperventilation.", "notes": "Branch-neutral once resuscitation reasonably maintained. Post-ROSC 12-lead shows Osborn (J) waves — classic hypothermia finding."},
            {"name": "Stage 4: Stabilised / Post-ROSC", "vitals": "HR 95 sinus, BP 78/50, SpO2 96%, T 35°C, GCS 6", "condition": "Stabilising; rewarming continuing; haemodynamics improving.", "expected": "TTM monitoring with an explicit temperature target stated. Actively avoids hyperthermia (>37.5°C). PICU transfer with rewarming plan handed over. Neurology referral for prognostication. Structured SBAR handover.", "notes": "End scenario."},
            {"name": "Stage 5: Rescue (capped)", "vitals": "Refractory VF or asystole despite Stage 2 management.", "condition": "Arrest persists despite adequate shocks, drugs and rewarming attempts.", "expected": "Continue CPR and active rewarming. ECMO referral/activation. Amiodarone per algorithm. Review reversible causes (4H4T).", "notes": "Automatic −10 Domain I penalty. Capped 45 s per the Shared Architecture, then scripted stabilisation."}
        ],
        "patient": {"name": "Vaibhavi", "mrn": "5501", "gender": "Female", "age": "12 years", "dob": "15-06-2014", "height": "145 cm", "weight": "30 kg", "cc": "Unresponsive after submersion", "hpi": "~10 min submersion in open lake water, ~15 min total CPR, two rounds of Adrenaline by EMS, T 33°C, no shock delivered yet.", "pmh": "None. Previously fit and well.", "psh": "None.", "meds": "Nil.", "allergies": "None.", "family": "Father (pulled her from the water) at the bedside, distressed. Mother not present."},
        "actors": "EMS crew (verbatim handover, then step back). Father — 'Is she going to be okay? I only looked away for a second…' Escalates emotionally if ignored >60 s; never obstructs. If asked submersion time: 'Maybe 10 minutes, I am not sure, it felt like forever.' Nurse prompts: 0:30 roles · 3:00 'Should I draw up another Adrenaline?' (no — shock first) · 5:00 rewarming escalation · 7:00 PICU bed.", "equipment": "Defibrillator + paediatric pads. Warming: forced-air blanket, warm IV fluid prop, warm humidified O2. Wet clothing moulage, towels. Monitor pre-loaded fine VF → sinus at ROSC; 12-lead with Osborn waves. Printed resource-answer cards. Crash cart."
    },
    {
        "id": "B2",
        "round": "BLS CASES",
        "title": "Torsades Arrest — Congenital Long QT Syndrome",
        "summary": "11-year-old female (30 kg) with known Long QT syndrome (non-compliant with Propranolol for a month) collapses at school — pulseless Torsades de Pointes. The team must name Torsades (not just 'VT'), give Magnesium sulphate 50 mg/kg alongside unsynchronised defibrillation 2 J/kg, correct K 2.5 / iCa 0.8, and explicitly AVOID Amiodarone (worsens QT).",
        "background": "Varsha, 11 years old, 30 kg. Known Long QT syndrome on beta-blockers, but stopped taking them for about a month. Collapsed suddenly at school, no pulse — teacher started CPR straight away (~4 min so far). Monitor shows an odd twisting rhythm. No IV access yet.",
        "expanded_history": "Diagnosed LQTS at age 6 after a similar collapse; usually on Propranolol, poor compliance for the last month per mother. No exertion or loud-noise trigger — she was sitting in class. Family history: maternal uncle died suddenly at 19, cause never confirmed. Two younger siblings not yet screened. Mother present, extremely anxious, asks repeatedly whether this happened because she 'kept forgetting her tablets.'",
        "stages": [
            {"name": "Stage 1: Initial Assessment (0:00–2:00)", "vitals": "HR: unrecordable. Rhythm: sinusoidal twisting polymorphic VT (Torsades). RR 0, BP unrecordable, unresponsive", "condition": "Pulseless arrest. Monitor shows twisting-around-baseline pattern — classic Torsades de Pointes, not monomorphic VT. Teacher's CPR in progress on arrival.", "expected": "Recognises arrest; continues high-quality CPR. Correctly identifies TORSADES DE POINTES, not simply 'VT'. Gains IV/IO access promptly. Requests urgent electrolytes (K, iCa, Mg). Prepares for UNSYNCHRONISED defibrillation — pulseless Torsades is shockable.", "notes": "Team names Torsades specifically and prepares Magnesium alongside the first shock → good trajectory. LQTS history not connected to the rhythm → nurse escalates at 3:00."},
            {"name": "Stage 2: Refractory Torsades (2:00–5:00)", "vitals": "Rhythm: Torsades persists / recurs intermittently. K 2.5 (on request), iCa 0.8 (on request)", "condition": "Refractory to shock alone without electrolyte correction — hypokalaemia and hypocalcaemia are sustaining the rhythm.", "expected": "MAGNESIUM SULPHATE 50 mg/kg IV (1.5 g) — the critical, most commonly missed intervention. Unsynchronised defibrillation 2 J/kg. Potassium replacement: KCl 0.5 mmol/kg IV. Explicitly AVOIDS Amiodarone — it further prolongs QT and worsens Torsades.", "notes": "Magnesium early alongside shocking → ROSC ~5:30. Magnesium omitted, or Amiodarone given → refractory, routes toward Stage 5."},
            {"name": "Stage 3: ROSC (≈5:30)", "vitals": "HR 110 sinus, BP 88/54, GCS 6 improving", "condition": "ROSC, scripted once adequate Magnesium + shock + K correction delivered.", "expected": "Recognises ROSC promptly, stops CPR, confirms pulse. Airway if needed: Ketamine + Rocuronium, avoiding further QT-prolonging agents. Continues Magnesium infusion and potassium correction. Actively excludes QT-prolonging drugs (Ondansetron, Metoclopramide) from the post-ROSC plan.", "notes": "Fires for every team once resuscitation reasonably maintained — branch-neutral."},
            {"name": "Stage 4: Stabilised / Post-ROSC", "vitals": "HR 98 sinus, BP 96/60, QTc still prolonged on repeat ECG", "condition": "Stabilising; QT remains prolonged, ongoing correction required.", "expected": "QTc monitoring continued explicitly. Magnesium infusion maintained, K target documented. Cardiology referral for ICD assessment discussed. Family screening for LQTS mentioned. Structured SBAR handover.", "notes": "End scenario."},
            {"name": "Stage 5: Rescue (capped)", "vitals": "Refractory Torsades despite Stage 2 management.", "condition": "Arrest persists — Magnesium omitted or Amiodarone given, sustaining the arrhythmia.", "expected": "Continues CPR. Gives Magnesium NOW if not already given. Corrects potassium. Avoids further QT-prolonging drugs. Continues defibrillation.", "notes": "Automatic −10 Domain I penalty. Capped 45 s, then scripted stabilisation."}
        ],
        "patient": {"name": "Varsha", "mrn": "5525", "gender": "Female", "age": "11 years", "dob": "05-11-2014", "height": "140 cm", "weight": "30 kg", "cc": "Sudden collapse at school, pulseless", "hpi": "Known LQTS, off Propranolol ~1 month. Sudden collapse while sitting in class. Teacher CPR within a minute, ~4 min so far.", "pmh": "Long QT syndrome diagnosed age 6 after a similar collapse.", "psh": "None.", "meds": "Propranolol — non-compliant for the last month.", "allergies": "None.", "family": "Maternal uncle died suddenly at 19. Two younger siblings unscreened. Mother present, guilt-ridden and anxious."},
        "actors": "Mother — 'Is this my fault? I kept forgetting her tablets…' Guilt-driven anxiety escalates if not reassured; never obstructs. If asked about triggers: 'No, she was just sitting in class, nothing happened, she just collapsed.' Nurse prompts: 0:30 roles · 3:00 'Should I draw up the Amiodarone?' (NO — Magnesium first) · 5:00 mother's guilt question · 7:00 Cardiology.", "equipment": "Defibrillator + paediatric pads. Monitor displaying Torsades → sinus at ROSC. Prop: old ECG printout QTc 480 ms (on request). Magnesium sulphate and KCl props. Printed resource-answer cards. Crash cart."
    },
    {
        "id": "B3",
        "round": "BLS CASES",
        "title": "Electrocution — VF Arrest with Hyperkalaemia",
        "summary": "13-year-old male (30 kg) contacted an overhead power line via a kite string — coarse VF arrest with entry/exit burns (right hand, left foot). The team must defibrillate 2 J/kg immediately, then recognise rhabdomyolysis-driven hyperkalaemia (K 7.5) as the cause of refractory VF: Calcium gluconate is the critical missed step, plus bicarbonate and insulin/dextrose. Succinylcholine is contraindicated for RSI.",
        "background": "Arush, 13 years old, 30 kg. Flying a kite near overhead power lines, got a big shock — thrown backward. Bystander CPR for ~3 minutes. Burns on right hand and left foot. No pulse, not breathing.",
        "expanded_history": "High-voltage overhead line contact via kite string; thrown ~2 metres. Entry wound right hand, exit wound left foot; clothing singed. Fire brigade made the scene safe before bystanders started CPR. Near a school playground — several distressed child witnesses. Father on his way, not yet arrived. No PMH, allergies or medications.",
        "stages": [
            {"name": "Stage 1: Initial Assessment (0:00–2:00)", "vitals": "HR: unrecordable (coarse VF), RR 0, apnoeic. Burns: entry right hand, exit left foot. K: pending", "condition": "Pulseless VF arrest immediately post-electrocution. Bystander CPR in progress on arrival.", "expected": "High-quality CPR commenced immediately. Rhythm identified: coarse VF — shockable. IMMEDIATE defibrillation 2 J/kg — not delayed for drug administration. Entry/exit burns identified; plans secondary survey for blast injury.", "notes": "Immediate defibrillation → Stage 2. Defibrillation delayed for drug preparation → nurse escalates at 3:00."},
            {"name": "Stage 2: Refractory VF — Hyperkalaemia Revealed (2:00–5:00)", "vitals": "Rhythm: coarse VF persists. K 7.5 (on request) — rhabdomyolysis-driven", "condition": "VF refractory to shock alone — severe hyperkalaemia from rhabdomyolysis is sustaining the arrhythmia.", "expected": "CALCIUM GLUCONATE 50 mg/kg IV — cardiac membrane stabilisation, the critical action. Sodium bicarbonate 1 mEq/kg IV — shifts potassium intracellularly. Insulin 0.1 U/kg + Dextrose 10% 2 mL/kg. Continued CPR and defibrillation each cycle.", "notes": "Calcium early → ROSC ~5:30 with a reasonable trajectory. Calcium omitted → refractory, routes toward Stage 5."},
            {"name": "Stage 3: ROSC (≈5:30)", "vitals": "HR 100 sinus, BP 84/50, K improving to 5.8", "condition": "ROSC, scripted once Calcium/Bicarb/CPR adequately delivered.", "expected": "Recognises ROSC promptly, stops CPR, confirms pulse. RSI with Ketamine + Rocuronium — SUCCINYLCHOLINE AVOIDED (worsens hyperkalaemia). Burns assessment: BSA estimated, Parkland formula initiated. Urinary catheter — myoglobinuria expected (cola-coloured urine).", "notes": "Fires for every team once resuscitation reasonably maintained — branch-neutral."},
            {"name": "Stage 4: Stabilised / Post-ROSC", "vitals": "HR 105 sinus, BP 92/58. Urine: dark, cola-coloured", "condition": "Stabilising; rhabdomyolysis and burns management ongoing.", "expected": "Ongoing K monitoring; Kayexalate or renal replacement planned if K stays elevated. Aggressive IV fluids for rhabdomyolysis, alkalinised urine target. Burns surgery referral and wound plan. Structured SBAR handover with K trend and fluid balance.", "notes": "End scenario."},
            {"name": "Stage 5: Rescue (capped)", "vitals": "Refractory VF, K remaining ≥7.5.", "condition": "Arrest persists — Calcium omitted or inadequate potassium-shifting therapy.", "expected": "Continues CPR. Gives Calcium gluconate NOW if not already given. Bicarbonate and Insulin/Dextrose. Continues defibrillation.", "notes": "Automatic −10 Domain I penalty. Capped 45 s, then scripted stabilisation."}
        ],
        "patient": {"name": "Arush", "mrn": "5526", "gender": "Male", "age": "13 years", "dob": "20-08-2013", "height": "150 cm", "weight": "30 kg", "cc": "Electrocution — pulseless, entry/exit burns", "hpi": "Kite string contacted overhead power line; thrown backward ~2 m. Bystander CPR ~3 min after fire brigade secured the scene.", "pmh": "None.", "psh": "None.", "meds": "Nil.", "allergies": "None.", "family": "Father on his way, not yet arrived. Child witnesses distressed."},
        "actors": "Bystander/father arriving — 'I just turned my back for a second, he was flying his kite…' Escalates emotionally if ignored; confirms kite-string/power-line mechanism if asked. Nurse prompts: 0:30 extra hands · 3:00 'His potassium's back — 7.5. Do you want calcium?' · 5:00 'His urine looks like cola…' · 7:00 burns/surgery call.", "equipment": "Defibrillator + paediatric pads. Burns dressing kit; moulage entry/exit wounds (right hand, left foot). Prop: kite string. Monitor pre-loaded coarse VF → sinus; ECG strip with peaked T waves. Urinary catheter prop with dark cola-coloured fluid. Printed resource-answer cards. Crash cart."
    },
    {
        "id": "B4",
        "round": "BLS CASES",
        "title": "Refractory Status Epilepticus with Raised ICP",
        "summary": "12-year-old female (30 kg) seizing 25 minutes despite Lorazepam and Levetiracetam en route, with shallow irregular breathing. Hidden inside the seizure: unequal pupils (R 5 mm > L 3 mm) and hypertension — raised ICP evolving toward a scripted herniation event at 5:30. The team must choose a 2nd-line non-benzodiazepine AED, give osmotherapy, position head-up, RSI with Ketamine (never Propofol), and call neurosurgery early.",
        "background": "Sita, 12 years old, 30 kg. Seizing for the last 25 minutes. Ambulance gave Lorazepam then Levetiracetam en route — neither worked. Still convulsing, breathing shallow and irregular.",
        "expanded_history": "No prior epilepsy. Two days of headache and vomiting before today; mother thought it was viral. No fever this morning, no trauma. Mother mentions Sita was 'more sleepy than usual' yesterday evening and once complained of blurry vision. No rash. Vaccinated. No travel.",
        "stages": [
            {"name": "Stage 1: Initial Assessment (0:00–2:00)", "vitals": "HR 130, RR shallow/irregular, BP 148/90 (hypertensive), Pupils R 5 mm > L 3 mm unequal, GCS 6 actively seizing", "condition": "Ongoing convulsive status epilepticus; airway at risk; early raised-ICP signs present but easy to miss mid-seizure.", "expected": "Immediate BVM ventilation for inadequate breathing. Recognises unequal pupils as raised ICP, not just the seizure. Identifies early Cushing triad: hypertension with relatively slow heart rate for the situation. Correct 2nd-line AED — Phenobarbitone 20 mg/kg or Valproate 40 mg/kg (NOT a third benzodiazepine).", "notes": "BVM + correct 2nd-line AED early → better trajectory. Another benzodiazepine, or airway delayed → nurse escalates at 3:00."},
            {"name": "Stage 2: Osmotherapy Window (2:00–5:00)", "vitals": "Seizure ongoing or terminating. BP rising further, HR relatively falling (early Cushing)", "condition": "Raised intracranial pressure evolving; window to intervene before herniation.", "expected": "3% NaCl 3 mL/kg IV or Mannitol 0.5 g/kg IV. Head elevated 30°, neutral neck. RSI with Ketamine + Rocuronium — PROPOFOL AVOIDED (drops MAP, worsens ICP). Early neurosurgery referral, before herniation.", "notes": "Osmotherapy + positioning + correct RSI + early neurosurgery call → team enters Stage 3 in a stronger position. Omissions compound the severity of the scripted event."},
            {"name": "Stage 3: Herniation Event (≈5:30)", "vitals": "HR 60 (bradycardia), BP 160/100, pupils bilateral fixed and dilated, GCS 3", "condition": "Acute herniation event — scripted, fires for every team regardless of prior performance. Poor Stage 2 management means arriving here in worse condition, not a different scenario.", "expected": "Recognises and NAMES herniation aloud. Escalates osmotherapy: additional 3% NaCl 5 mL/kg or Mannitol 1 g/kg. Brief hyperventilation (PCO2 target 30–35) as a bridge, not prolonged. Neurosurgery emergency activation and urgent imaging plan.", "notes": "Scripted and unpreventable — do not score a prevention pathway."},
            {"name": "Stage 4: Stabilised", "vitals": "Seizure electrographically controlled. Pupils improving. HR/BP normalising", "condition": "Stabilising; ICP crisis responding to escalated therapy.", "expected": "Continuous AED infusion planned. Temperature managed — paracetamol; hyperthermia worsens seizure and ICP. PICU transfer with neurosurgery on standby. ICP monitoring plan communicated. Structured SBAR handover.", "notes": "End scenario."},
            {"name": "Stage 5: Rescue (capped)", "vitals": "Uncontrolled ICP crisis / seizure refractory to Stage 2 measures.", "condition": "Persistent raised ICP and/or ongoing seizure activity.", "expected": "Continues airway support. Gives osmotherapy NOW if not already given. Urgent neurosurgery escalation. Considers further AED.", "notes": "Automatic −10 Domain I penalty. Capped 45 s, then scripted stabilisation."}
        ],
        "patient": {"name": "Sita", "mrn": "5527", "gender": "Female", "age": "12 years", "dob": "15-05-2014", "height": "148 cm", "weight": "30 kg", "cc": "Status epilepticus 25 min, failed 1st/2nd line en route", "hpi": "Two days of headache and vomiting, then seizures. Lorazepam + Levetiracetam en route without effect. Breathing shallow and irregular.", "pmh": "No prior epilepsy. Previously well.", "psh": "None.", "meds": "Received Lorazepam + Levetiracetam en route.", "allergies": "None.", "family": "Mother present — 'She said her head hurt for two days, I thought it was just a bug…'"},
        "actors": "Mother — provides expanded history if asked; escalates emotionally if ignored, never obstructs. Nurse prompts: 0:30 'Do you want me to bag her?' · 3:00 'Should I draw up more Lorazepam?' (NO — 2nd-line non-benzo AED) · 5:00 'Her right pupil looks different to her left…' · 7:00 neurosurgery call.", "equipment": "BVM, RSI kit (Ketamine, Rocuronium). Manikin/monitor capable of pupil asymmetry. 3% NaCl / Mannitol infusion props. Monitor trend consistent with Cushing triad at the herniation event. Printed resource-answer cards. Crash cart."
    },
    {
        "id": "P1",
        "round": "PALS CASES",
        "title": "Cardiogenic Shock — Fulminant Myocarditis",
        "summary": "1-year-old (10 kg), three days of flu-like illness, now pale, breathless, cool with a gallop rhythm and hepatomegaly — fulminant myocarditis in cardiogenic shock. The fluid bolus is the trap. At ~5:30 the monitor VT becomes pulseless: it still needs SYNCHRONISED cardioversion 0.5 J/kg (organised rhythm), not unsynchronised 2 J/kg. Inotrope first, Ketamine-only RSI.",
        "background": "Sita, 1 year old, 10 kg. Three days of what looked like flu — fever, unsettled, off her feeds. Today very pale and breathless; lips looked blue to mum. Cool to touch, heart rate very fast.",
        "expanded_history": "Three days of URTI-like illness (fever, malaise); no vomiting or diarrhoea. Previously fit and well, no cardiac or family history. Mother: much less active today, floppier than normal, briefly less responsive once. No rash. Vaccinated. No toxin/drug ingestion.",
        "stages": [
            {"name": "Stage 1: Initial Assessment (0:00–2:00)", "vitals": "HR 170, BP 78/50, CRT 4 s, GCS 13. Exam: cold peripheries, gallop rhythm, hepatomegaly, crackles", "condition": "Circulatory failure consistent with cardiogenic shock — not septic (no current fever, insidious onset, low-output signs) and not hypovolaemic.", "expected": "Rapid ABCDE identifying circulatory failure. Identifies CARDIOGENIC, not septic or hypovolaemic, shock. Requests 12-lead ECG, point-of-care echo, troponin, BNP.", "notes": "Reflexive large fluid bolus without reassessment → pulmonary oedema worsens, faster deterioration. Correct restraint → better trajectory. Nurse prompt 0:30: 'Do you want a fluid bolus running, doctor?' — correct answer is caution/avoid."},
            {"name": "Stage 2: Inotrope Window (2:00–5:00)", "vitals": "Echo (if requested): EF ~15%. Monitor: organised wide-complex tachycardia (VT), rate 220, not yet pulseless", "condition": "Deteriorating cardiogenic shock; VT appears on the monitor as an organised rhythm.", "expected": "ADRENALINE INOTROPE 0.05–0.1 mcg/kg/min — NOT a fluid bolus as first-line. Actively avoids/halts fluid boluses on recognising cardiogenic shock. If intubation needed: Ketamine + Rocuronium — NOT Propofol or Midazolam. Recognises VT as an organised rhythm requiring a synchronised approach if unstable/pulseless.", "notes": "Inotrope-first → team enters the scripted VT event from a stronger position. Fluid bolus given → pulmonary oedema worsens into Stage 3."},
            {"name": "Stage 3: Pulseless VT (≈5:30)", "vitals": "HR 220 (VT on monitor), pulse absent, BP unrecordable", "condition": "Scripted event: VT becomes pulseless. Fires for every team ~5:30, branch-neutral occurrence.", "expected": "Recognises pulseless VT, calls for the defibrillator. Selects SYNCHRONISED mode, 0.5 J/kg — NOT unsynchronised 2 J/kg (the rhythm is organised). Resumes CPR immediately after the shock without a prolonged pause. Post-ROSC: continues inotrope targeting MAP >60, plans airway and PICU transfer.", "notes": "A team that avoided fluid and started the inotrope early arrives here in better condition."},
            {"name": "Stage 4: Stabilised / Post-ROSC", "vitals": "Post-shock: organised rhythm returns, BP improving on inotrope", "condition": "Stabilising on Adrenaline infusion.", "expected": "Inotrope titrated, MAP >60 target stated. Airway plan (Ketamine RSI) if not yet secured. Cardiology + PICU referral. Structured SBAR handover.", "notes": "End scenario."},
            {"name": "Stage 5: Rescue (capped)", "vitals": "Recurrent pulseless VT / unable to achieve ROSC.", "condition": "Refractory cardiogenic shock with recurrent arrhythmia.", "expected": "Continues CPR. Repeats synchronised or unsynchronised shock per rhythm. Escalates inotrope support. Considers ECMO referral for refractory cardiogenic shock.", "notes": "Automatic −10 Domain I penalty. Capped 45 s, then scripted stabilisation."}
        ],
        "patient": {"name": "Sita", "mrn": "5529", "gender": "Female", "age": "1 year", "dob": "20-10-2024", "height": "75 cm", "weight": "10 kg", "cc": "Pale, breathless, cool — 3 days flu-like illness", "hpi": "Three days fever/malaise, off feeds. Today pale, breathless, briefly less responsive, lips dusky.", "pmh": "Previously well.", "psh": "None.", "meds": "Nil.", "allergies": "None.", "family": "Mother present — 'She's never been this sick before, is her heart okay?'"},
        "actors": "Mother — frightened, provides history on request, escalates if ignored, never obstructs. Nurse prompts: 0:30 fluid bolus offer (trap) · 3:00 'Should I draw up another 20 mL/kg bolus?' · 5:00 intubation trolley · 7:00 PICU call.", "equipment": "Defibrillator with synchronised-mode capability. Manikin/monitor displaying wide-complex VT → pulseless → post-shock rhythm. Echo image/prop card (EF ~15%). Inotrope infusion pump prop. Printed resource-answer cards. Crash cart."
    },
    {
        "id": "P2",
        "round": "PALS CASES",
        "title": "Septic Shock — Post-Open Heart Surgery",
        "summary": "2-year-old (8 kg), three weeks post-ASD repair, drowsy and floppy with critical hypoglycaemia (glucose 28) and a wide-QRS tachycardia that is a trap — sinus tachycardia with rate-related aberrancy from hyponatraemia (Na 122), NOT VT. The team must give dextrose within minutes, draw cultures then start Meropenem + Vancomycin, find the infected sternotomy wound, and explicitly NOT cardiovert.",
        "background": "Manoj, 2 years old, 8 kg. Three weeks post-op from an ASD repair, recovering at home. Off-colour since yesterday, now drowsy and floppy. Blood sugar on the ambulance monitor read low.",
        "expanded_history": "ASD closure three weeks ago, uncomplicated recovery until yesterday. Not eating or drinking well for 24 h. No documented fever at home. No regular medications. Mother mentions the surgical wound looked 'a bit red' two days ago. Vomiting overnight; no diarrhoea; no rash.",
        "stages": [
            {"name": "Stage 1: Initial Assessment (0:00–2:00)", "vitals": "GCS 11 drowsy, HR 165, BP 76/40, CRT 4 s, Glucose 28 (critical). Monitor: wide-QRS tachycardia", "condition": "Septic shock, likely post-surgical wound source, presenting with critical hypoglycaemia and a rhythm that looks alarming but is a trap.", "expected": "Recognises septic shock within the first minute. Identifies CRITICAL HYPOGLYCAEMIA (28) and treats within ~60 s. Interprets the wide-QRS rhythm as sinus tachycardia from hyponatraemia — NOT VT, does NOT cardiovert. Identifies Na 122 on first bloods and plans symptomatic correction.", "notes": "Reflexive cardioversion of the wide-QRS rhythm is a critical, avoidable error — case deteriorates and nurse heavily escalates. Correct restraint + dextrose → better trajectory."},
            {"name": "Stage 2: Source Control (2:00–5:00)", "vitals": "Glucose rising post-treatment. Na 122. HR improving slightly", "condition": "Septic shock with an identified surgical-wound source and correctable metabolic derangements.", "expected": "Dextrose 10% 2 mL/kg promptly, within 2 minutes. Blood cultures drawn THEN broad-spectrum antibiotics (Meropenem + Vancomycin). Cardioversion explicitly NOT performed. Vasopressor (Noradrenaline or Adrenaline) initiated + 3% NaCl 2 mL/kg for Na 122.", "notes": "Correct sequence → stabilising trajectory into the scripted seizure. Delays in dextrose/antibiotics or a mistaken cardioversion → team arrives at Stage 3 in worse condition. Nurse 5:00: 'His wound dressing looks a bit soggy and red…' — examine the wound."},
            {"name": "Stage 3: Seizure Event (≈5:30)", "vitals": "GCS drops to 8 during seizure, HR 150, SpO2 88% (airway compromise)", "condition": "Scripted brief generalised seizure — fires for every team ~5:30, severity modulated by how well glucose/Na were corrected beforehand.", "expected": "Airway management: lateral position, suction, BVM if apnoeic. Benzodiazepine — Midazolam 0.15 mg/kg. Recognises the seizure as a consequence of hyponatraemia/hypoglycaemia and continues electrolyte correction.", "notes": "Scripted and unpreventable; severity reflects prior correction quality."},
            {"name": "Stage 4: Stabilised", "vitals": "GCS improving to 13, HR 120, BP 92/58 on vasopressor", "condition": "Stabilising on antibiotics, vasopressor, and electrolyte correction.", "expected": "Antibiotics continued; vasopressor titrated to MAP >60. Sodium trend monitored — correction not excessive (max 8 mEq/L in 24 h). Cardiology informed of a post-op child in septic shock. Structured SBAR handover for PICU.", "notes": "End scenario."},
            {"name": "Stage 5: Rescue (capped)", "vitals": "Ongoing shock / seizure refractory to Stage 2 measures.", "condition": "Persistent hypoglycaemia/hyponatraemia-driven instability.", "expected": "Continues resuscitation. Repeats dextrose / sodium correction. Escalates vasopressor. Urgent PICU escalation.", "notes": "Automatic −10 Domain I penalty. Capped 45 s, then scripted stabilisation."}
        ],
        "patient": {"name": "Manoj", "mrn": "5533", "gender": "Male", "age": "2 years", "dob": "10-03-2024", "height": "84 cm", "weight": "8 kg", "cc": "Drowsy and floppy, 3 weeks post-ASD repair", "hpi": "Off-colour since yesterday, poor intake 24 h, vomiting overnight. Ambulance glucose read low. Wound looked red two days ago.", "pmh": "ASD surgical closure 3 weeks ago, uncomplicated until now.", "psh": "ASD repair (sternotomy).", "meds": "None currently.", "allergies": "None.", "family": "Mother present — 'He was doing so well after his surgery, I don't understand what's happened…'"},
        "actors": "Mother — provides wound history/timeline if asked; escalates if ignored, never obstructs. Nurse prompts: 0:30 low glucose strip · 3:00 'That rhythm looks fast and wide, should I get the defibrillator?' (NO — sinus tach, not VT) · 5:00 soggy red wound dressing · 7:00 PICU/cardiology.", "equipment": "Point-of-care glucose meter, dextrose prep. Manikin/monitor displaying wide-QRS sinus tachycardia narrowing as rate/electrolytes correct. Surgical wound moulage (erythematous, post-sternotomy). Midazolam prop for the seizure. Printed resource-answer cards. Crash cart."
    },
    {
        "id": "P3",
        "round": "PALS CASES",
        "title": "Acute Severe Asthma with Tension Pneumothorax",
        "summary": "3-year-old known asthmatic (14 kg), worsening 2 days despite inhalers and home nebulisers — near-fatal asthma with a quietening chest. Mother is steroid-hesitant and needs empathy, not confrontation. Early IV Magnesium 50 mg/kg, Ketamine-only RSI; then a scripted right tension pneumothorax immediately post-intubation at 5:30 — needle decompression BEFORE any chest X-ray.",
        "background": "Arun, 3 years old, 14 kg. Known asthmatic, worse over the last two days despite inhalers. Home nebulisers with no improvement. Really struggling to breathe, single words only.",
        "expanded_history": "Asthma diagnosed at age 2, usually well controlled on a preventer; a few missed doses recently. Mild cough ~1 week, no clear viral trigger. No allergies. Father also asthmatic. Mother is hesitant about steroids — a family member had a difficult experience on long-term steroids; needs a careful, empathetic explanation, not confrontation.",
        "stages": [
            {"name": "Stage 1: Initial Assessment (0:00–2:00)", "vitals": "RR 44, SpO2 87% on air, HR 150. Exam: tripod positioning, accessory muscle use, chest becoming quiet", "condition": "Near-fatal asthma — rising work of breathing, falling SpO2 despite nebulisers already tried at home.", "expected": "Recognises near-fatal asthma. Navigates parental steroid refusal EMPATHETICALLY — explains life-threatening risk, documents if still refused. Initiates IV Magnesium sulphate 50 mg/kg over 20 minutes early. Identifies deteriorating respiratory failure requiring an intubation decision.", "notes": "Empathetic negotiation + early Magnesium → better trajectory. Dismissive handling of the refusal, or delayed Magnesium → nurse escalates at 3:00."},
            {"name": "Stage 2: Failing Ventilation (2:00–5:00)", "vitals": "SpO2 falling to 82% despite nebulisers/Magnesium. Patient tiring, becoming drowsy", "condition": "Progressing toward respiratory failure despite maximal medical therapy.", "expected": "IV Magnesium 50 mg/kg — correct dose calculated (0.7 g for 14 kg). RSI with Ketamine + Rocuronium — PROPOFOL AVOIDED (bronchoconstriction). Prepares for intubation as work of breathing fails.", "notes": "Correct RSI agent + timely intubation → team enters Stage 3 in reasonable shape. Propofol chosen, or intubation delayed → deterioration compounds."},
            {"name": "Stage 3: Post-Intubation Tension Pneumothorax (≈5:30)", "vitals": "SpO2 falling further, airway pressure rising, absent right-sided breath sounds, HR rising, BP falling", "condition": "Scripted new right tension pneumothorax immediately post-intubation. Fires for every team ~5:30; a team that intubated well on Ketamine reaches this in better shape.", "expected": "Recognises post-intubation deterioration as an acute event. NEEDLE DECOMPRESSION BEFORE requesting a CXR — right side, 2nd intercostal space, mid-clavicular line. Chest drain after needle decompression, reconnects ventilator. PEEP reduced post-decompression; ventilator optimised for asthma (long expiratory time).", "notes": "The bag feels stiff, sats drop, right chest silent — decompress first, image second."},
            {"name": "Stage 4: Stabilised", "vitals": "Breath sounds bilateral, restored. SpO2 improving to 94%. HR settling", "condition": "Stabilising after decompression and drain insertion.", "expected": "Post-pneumothorax reassessment: bilateral breath sounds, SpO2 improved. Magnesium infusion continued. Steroid now given if earlier refused. PICU referral with explicit ventilator settings. Structured SBAR handover.", "notes": "End scenario."},
            {"name": "Stage 5: Rescue (capped)", "vitals": "Ongoing hypoxia / tension physiology despite Stage 3 measures.", "condition": "Persistent respiratory failure with unresolved tension physiology.", "expected": "Continues decompression/drain troubleshooting. Escalates ventilation support. Urgent PICU escalation.", "notes": "Automatic −10 Domain I penalty. Capped 45 s, then scripted stabilisation."}
        ],
        "patient": {"name": "Arun", "mrn": "5531", "gender": "Male", "age": "3 years", "dob": "20-10-2022", "height": "92 cm", "weight": "14 kg", "cc": "Severe breathing difficulty, single words only", "hpi": "Worsening over 2 days despite inhalers and home nebulisers. Mild cough ~1 week.", "pmh": "Asthma diagnosed age 2; preventer inhaler with recent missed doses.", "psh": "None.", "meds": "Preventer inhaler (adherence lapses); home nebulisers today.", "allergies": "None.", "family": "Father also has asthma. Mother present, steroid-hesitant."},
        "actors": "Mother — 'I've seen what long-term steroids did to my brother, I don't want that for him.' Guarded but not obstructive; settles once given an empathetic short-course explanation. Nurse prompts: 0:30 steroid refusal relay · 3:00 Magnesium draw-up · 5:00 'Sats dropping fast and the bag feels really stiff…' · 7:00 PICU.", "equipment": "Nebuliser setup, IV Magnesium prop. RSI kit (Ketamine, Rocuronium). Needle decompression + chest drain kits. Manikin simulating unilateral absent breath sounds. Ventilator/BVM prop with adjustable resistance feel. Printed resource-answer cards. Crash cart."
    },
    {
        "id": "P4",
        "round": "PALS CASES",
        "title": "Refractory Septic Shock — Physiologically Difficult Airway",
        "summary": "3-year-old (15 kg) with sudden fever and a spreading non-blanching purpuric rash — meningococcal septicaemia in refractory septic shock, essentially peri-arrest. RESUSCITATE BEFORE YOU INTUBATE: vasopressor running and push-dose Epinephrine 10 mcg/kg drawn up at the bedside before any induction; Ketamine is the only safe RSI agent. Ceftriaxone after cultures, hydrocortisone 2 mg/kg, and a scripted peri-intubation crash at 5:30.",
        "background": "Maya, 3 years old, 15 kg. Sudden high fever and a rash that's not fading on pressure, started a few hours ago. Very floppy in the last 30 minutes, barely responding. Ambulance BP really low.",
        "expanded_history": "Previously well. Fever onset six hours ago; rapidly spreading non-blanching purpuric rash for ~2 h. One EMS fluid bolus with minimal response. Vaccinated incl. MenACWY per mother — breakthrough possible. If pressed: recent family travel to a region with a reported meningococcal outbreak. Younger sibling at home, well. Father anxious, asks whether the family will 'catch it.'",
        "stages": [
            {"name": "Stage 1: Initial Assessment (0:00–2:00)", "vitals": "BP 50/30, HR 175, Lactate 9.2, GCS 9, CRT 5 s. Rash: non-blanching purpuric, spreading", "condition": "Refractory septic shock — essentially peri-arrest, zero physiological reserve. One EMS bolus already given with minimal response.", "expected": "Recognises refractory septic shock as a peri-arrest state. Identifies the non-blanching purpuric rash as MENINGOCOCCAL SEPTICAEMIA. Decision: RESUSCITATE FIRST — vasopressor and fluid before any intubation attempt. Ceftriaxone immediately after blood cultures.", "notes": "Any RSI attempt before a vasopressor is running is a critical safety flag. Nurse 0:30: 'Do you want the intubation trolley ready?' — correct answer is NOT YET, resuscitate first."},
            {"name": "Stage 2: Resuscitate Before Intubate (2:00–5:00)", "vitals": "BP persistently low despite further fluid. Rash worsening. GCS falling further", "condition": "Persistent hypotension despite fluid; airway will need securing, but sequencing is critical.", "expected": "Vasopressor (Noradrenaline or Adrenaline) initiated BEFORE any RSI attempt, dose stated. RSI agent: KETAMINE ONLY — not Propofol, Midazolam, or Fentanyl. Push-dose Epinephrine 10 mcg/kg IV prepared at the bedside before intubation. Hydrocortisone 2 mg/kg IV — relative adrenal insufficiency / Waterhouse-Friderichsen.", "notes": "All four measures present → team enters the scripted crash from a stronger position. Missing push-dose Epinephrine especially worsens the 5:30 crash."},
            {"name": "Stage 3: Peri-Intubation Crash (≈5:30)", "vitals": "BP crashes to 30/15 or unrecordable immediately post-induction. HR spikes then may drop. SpO2 falling", "condition": "Scripted peri-intubation cardiovascular crash at the moment of induction — fires for every team ~5:30; worse without push-dose Epinephrine ready.", "expected": "IMMEDIATELY administers push-dose Epinephrine 10 mcg/kg IV — delay beyond 30 s scored down. Reduces PEEP / ventilation pressure. Fluid push. Escalates vasopressor infusion.", "notes": "Scripted and unpreventable; severity is the consequence of Stage 2 preparation."},
            {"name": "Stage 4: Stabilised", "vitals": "BP 78/48 on increased vasopressor, HR 140. Rash static, not spreading", "condition": "Stabilising on escalated vasopressor support.", "expected": "Antibiotics continued; vasopressors titrated to MAP >60. Meningococcal contacts identified and notified — public health duty. PICU transfer plan with explicit vasopressor dose and rate. Structured SBAR handover.", "notes": "End scenario."},
            {"name": "Stage 5: Rescue (capped)", "vitals": "Cardiac arrest / unable to recover from the peri-intubation crash.", "condition": "Refractory shock progressing to arrest.", "expected": "CPR. Repeats push-dose Epinephrine / starts Adrenaline infusion. Continues vasopressor escalation. Urgent PICU / ECMO discussion.", "notes": "Automatic −10 Domain I penalty. Capped 45 s, then scripted stabilisation."}
        ],
        "patient": {"name": "Maya", "mrn": "5532", "gender": "Female", "age": "3 years", "dob": "15-08-2022", "height": "90 cm", "weight": "15 kg", "cc": "Sudden fever, purpuric rash, floppy and barely responding", "hpi": "Fever 6 h, non-blanching purpuric rash spreading over 2 h, one EMS bolus with minimal response.", "pmh": "Previously well. Vaccinated incl. MenACWY.", "psh": "None.", "meds": "Nil.", "allergies": "None.", "family": "Father present — anxious about household contacts. Younger sibling at home, well."},
        "actors": "Father — 'Is she going to be okay? Are the rest of us going to catch this?' Needs brief prophylaxis reassurance without delaying care; escalates if ignored, never obstructs. Nurse prompts: 0:30 intubation trolley (not yet) · 3:00 persistent hypotension · 5:00 at induction 'Should I push the drugs?' · 7:00 PICU/ID call.", "equipment": "Vasopressor infusion pump prop, push-dose Epinephrine prefilled syringe prop. RSI kit with Ketamine only visible. Non-blanching purpuric rash moulage. Hydrocortisone prop. Printed resource-answer cards. Crash cart."
    },
    {
        "id": "SF1",
        "round": "SEMI-FINALS",
        "title": "The Quiet Head — Dual Paediatric Trauma",
        "summary": "Two siblings arrive simultaneously after a high-speed RTA. Ananya Rao (7y, 24kg) is initially talking — the classic lucid interval of an expanding right temporal extradural haematoma. Vihaana Rao (18mo, 10kg) cries loudly with a femoral fracture and trace FAST finding, but is haemodynamically stable. The team must simultaneously manage both patients, resist false reassurance from Ananya's initial calm presentation, resist anchoring on an outside CT report that is genuinely normal — taken 90 minutes post-injury, before the haematoma was visible — and escalate the neurological emergency in time.",
        "background": "You are the paediatric trauma team on duty. Two siblings from a high-speed RTA arrive simultaneously. Ananya (7y) had brief LOC, vomited once, and is increasingly sleepy. Vihaana (18mo) has no LOC but cries persistently with a deformed right leg. IV access in both pre-hospital. CT brain done outside — report with father.",
        "expanded_history": "Ananya Rao (7y, 24kg): Restrained rear-seat passenger. Head struck right window. Brief LOC then recovery. Hidden diagnosis: right temporal extradural haematoma, temporal bone fracture, mild pulmonary contusion, closed distal radius fracture.\n\nVihaana Rao (18mo, 10kg): Child restraint partially detached. No LOC. Persistent crying. Hidden diagnosis: closed right femoral shaft fracture, trace perihepatic free fluid on FAST.",
        "stages": [
            {"name": "Stage 1: Act I — Golden Minute (0:00–4:00)", "vitals": "Ananya: HR 118, RR 28, BP 112/72, SpO2 97% NRB, GCS 13, Pupils 3mm equal, Vihaana: HR 168, RR 42, BP 96/60, SpO2 96%", "condition": "Ananya: Lying quietly. Responds slowly. Blood over right temporal scalp. Right wrist swollen. Lucid interval — no deterioration yet.\n\nVihaana: Crying loudly. Right leg held still. Seatbelt bruising abdomen. FAST: minimal perihepatic free fluid — haemodynamically stable.", "expected": "Designate team leader · Allocate sub-group to each child · Simultaneous ABCDE on both · Monitors on both · C-spine maintained for Ananya · Baseline GCS and pupils documented for Ananya · Trauma bloods for both · FAST for Vihaana · Recognise mechanism-based red flags from handover (Ananya: LOC, vomiting, increasing sleepiness)", "notes": "Deliberate calm opening — the lucid interval must tempt false reassurance. No deterioration in this window. Nurse prompt 3:00 if bloods not sent. Vihaana's FAST: minimal free fluid = calibrated response required (surgical consult + serial exam, not theatre)."},
            {"name": "Stage 2: Reassessment — GCS Trend + Parents Arrive (4:00–5:30)", "vitals": "Ananya: HR 100, BP 132/82, GCS 11 (E2V4M5), Pupils R4mm sluggish L3mm, Vihaana: HR 148, BP 96/60, SpO2 96%", "condition": "Ananya: GCS 13→11. Right pupil 4mm sluggish — anisocoria developing. BP rises (early Cushing). Father arrives 5:00 at Ananya's bay with outside CT report and X-rays.\n\nVihaana: Stable after analgesia and fluids. Mother arrives 5:00 at Vihaana's bay.", "expected": "Recognise GCS trend (not a single reading) · Identify new anisocoria · Read the outside CT report; recognise it is genuinely NORMAL (taken 90 min post-injury — an early EDH may not yet be visible) · Refuse to anchor on the normal report: the GCS and pupil trend is the diagnosis · Activate neurosurgery immediately · Prepare RSI airway equipment · Calibrated response to Vihaana's trace FAST: surgical consult + serial exam", "notes": "Nurse prompt 4:30: 'Who needs the USG machine first — Ananya or Vihaana?' One machine only — forces explicit resource decision. Father (5:00): 'She was talking all the way. The scan was normal. Why is she becoming sleepy?' The report truthfully reads normal — the trap is anchoring on it instead of the evolving GCS and pupils. Debrief reveal: an EDH imaged very early may genuinely be absent on the first scan."},
            {"name": "Stage 3: Scripted Acute Event — Seizure (5:30–8:00)", "vitals": "Ananya: HR 48, BP 138/90, GCS 6, Pupils R6mm fixed L3mm reactive, Irregular respirations", "condition": "Ananya seizes at 5:30 — scripted and unpreventable for both teams.\nBranch A (correct): RSI + ICP management → Stage 3 impending herniation (reversible).\nBranch B (delayed): No airway or ICP treatment → Stage 3B established herniation (GCS 3, apnoeic).", "expected": "Branch A (Correct): Airway · O2 · Suction · Benzodiazepine · Recognise Cushing response (bradycardia HR 48, rising BP, irregular respirations, fixed pupil) · RSI intubation with C-spine · 3% hypertonic saline 3–5 mL/kg · ETCO2 35–40 post-intubation · Head elevation 30° · Activate neurosurgery.\nBranch B (Incorrect): Delay intubation OR miss raised ICP signs → progresses to Stage 3B", "notes": "Seizure fires regardless of team performance. Do not score a prevention pathway. Father at Ananya's bay: 'She was talking 10 minutes ago — why does she need a breathing tube?' One team member communicates; leader does not stop managing patient. Cushing bradycardia (HR 48) is correct physiology, not an operator error. Stage 5 cap: 60 seconds, scripted stabilisation to Stage 3B, −10 penalty."},
            {"name": "Stage 4: Stabilised / Endgame (8:00–15:00)", "vitals": "Ananya: HR 91, BP 132/72, SpO2 97%, Intubated, Pupils R fixed L reactive, Vihaana: HR 138, BP 96/58, Splinted", "condition": "Ananya: Intubated, AED given, neuroprotection performed. Right pupil still fixed — requires surgical evacuation (not a failure of medical management).\n\nVihaana: Pain controlled, femur splinted, surgical review arranged.", "expected": "Confirm ETT placement · Continue sedation, maintain normocapnia · Prepare urgent neurosurgery transfer for Ananya · Arrange surgical review for Vihaana's abdominal injury and femur · Clear SBAR handover for both children", "notes": "Fixed pupil persists despite correct medical stabilisation — full recovery requires surgical evacuation. Do not let teams read unchanged pupil as failure. Tiebreaker: 'One CT scanner, one theatre. Justify the order of use for both children, and state what would change that order.'"}
        ],
        "patient": {"name": "Ananya Rao + Vihaana Rao", "mrn": "SF-01 / SF-02", "gender": "Female + Male", "age": "7 years + 18 months", "dob": "High-speed RTA — restrained rear-seat passengers", "height": "Ananya ~24 kg · Vihaana ~10 kg", "weight": "Dual-patient round — 15:00 hard stop", "cc": "Head injury (Ananya) + multiple trauma (Vihaana)", "hpi": "Ananya: Head struck right window, brief LOC, vomited once, increasingly sleepy. Vihaana: Child restraint partially detached, persistent crying, right leg deformed.", "pmh": "Both previously well.", "psh": "None.", "meds": "Nil.", "allergies": "None.", "family": "Father arrives 5:00 at Ananya's bay with outside CT report. Mother arrives 5:00 at Vihaana's bay."},
        "actors": "Ambulance Crews 1 and 2 (hand over, step back). One circulating trauma nurse initially. Father with CT report (Ananya's bay, 5:00). Mother (Vihaana's bay, 5:00).",
        "equipment": "Two multiparameter monitors. ONE ultrasound machine (contention is deliberate). Paediatric C-spine collars. Airway trolley + RSI tray. 3% hypertonic saline. AED vial/label. Printed outside CT report prop (normal, timed 90 min post-injury) + X-rays. Toddler femur splint. Paediatric BVM. Crash cart."
    },
    {
        "id": "SF2",
        "round": "SEMI-FINALS",
        "title": "The Loud Wound — Dual Paediatric Trauma",
        "summary": "The deliberate mirror image of SF1. Meera Iyer (8y, 26kg) arrives quiet — a grade IV splenic laceration with evolving haemorrhagic shock, misread by paramedics as stable. Diya Iyer (3y, 14kg) arrives with a visibly blood-soaked scalp dressing requiring two changes en route — but is neurologically and haemodynamically intact. The team must resist the pull of the more dramatic presentation, correctly interpret a positive FAST in an unstable child as an indication for surgery not CT, and activate massive transfusion.",
        "background": "You are the paediatric trauma team. Two siblings from a high-speed T-bone collision arrive simultaneously. Diya (3y) has a visibly blood-soaked scalp dressing. Meera (8y) is quiet — crew 'weren't as worried about her.' Both were restrained rear-seat passengers.",
        "expanded_history": "Meera Iyer (8y, 26kg): Seatbelt mark across abdomen. Quiet throughout transport. Crew misread as stable. Hidden: grade IV splenic laceration, progressive haemoperitoneum, Class II→III haemorrhagic shock, closed right femoral shaft fracture.\n\nDiya Iyer (3y, 14kg): Head struck side window. No LOC. Scalp bled briskly, dressing changed twice en route. Hidden: scalp laceration with subgaleal haematoma — NO skull fracture, NO intracranial injury, NO other injuries.",
        "stages": [
            {"name": "Stage 1: Act I — Both Arrive (0:00–4:00)", "vitals": "Meera: HR 138, RR 30, BP 100/68, SpO2 98%, CRT 3s, Pulse pressure narrow, Diya: HR 130, RR 28, BP 94/60, SpO2 99%, GCS 15", "condition": "Meera: Quiet, seatbelt bruising abdomen and right thigh. Pale, cool peripheries, narrowing pulse pressure — compensated shock.\n\nDiya: Crying loudly. Scalp dressing visibly blood-soaked. Large boggy swelling underneath. Alert, moving all limbs normally.", "expected": "Designate leader · Resist initial pull toward Diya (visible blood is disproportionate to clinical significance) · Allocate sub-group to each child · Recognise Meera's narrow pulse pressure as compensated shock · Two large-bore IV access on Meera · Type-and-crossmatch urgently for both · FAST for Meera · Firm direct pressure/pressure dressing for Diya's scalp · Analgesia for both", "notes": "The volume of visible blood from Diya is disproportionate to clinical significance — the scalp is highly vascular. FAST is not indicated for Diya (no abdominal mechanism) — requesting it is a resource error, not a safety error. Nurse prompt 3:00: 'Shall I send trauma bloods and crossmatch for both children?'"},
            {"name": "Stage 2: Reassessment — Shock Trend + Parents Arrive (4:00–5:30)", "vitals": "Meera: HR 156, RR 34, BP 92/60, SpO2 97%, CRT 3–4s, GCS 14, FAST positive, Diya: HR 116, RR 24, BP 98/62, GCS 15", "condition": "Meera: Rising HR, narrowing pulse pressure. GCS 14 — becoming drowsy. Abdomen more distended and tense. Reduced pain reporting = altered sensorium, NOT improvement. FAST: significant free fluid in Morrison's pouch, splenorenal recess and pelvis.\n\nDiya: Bleeding controlled by pressure. At 4:30 she vomits once and becomes briefly drowsy, then settles — GCS 15 on reassessment. Boggy swelling stable in size. Father arrives 5:00 at Meera's bay — initially reassured because she looks quieter.", "expected": "Recognise falling pulse pressure + rising HR as progression · Do NOT mistake Meera's reduced pain reporting for improvement · Positive FAST in haemodynamically unstable child = urgent surgery/IR, NOT CT · Escalate from crystalloid to blood products · Activate MTP (PRBC:FFP:platelets 1:1:1) · TXA 15 mg/kg over ~10 min · Active warming · Urgent surgical review · Explicitly apply a paediatric head-injury decision rule to Diya's 4:30 vomit and drowsy episode — state the conclusion aloud (neither reflex CT nor unreasoned dismissal) · Continue Diya's wound management without diverting team resources from Meera", "notes": "Nurse prompt 4:30: 'Diya's dressing has soaked through again — do you want more gauze, or should I get help?' Deliberate pull toward the louder patient. CT request for Meera is the trap — answer delivered neutrally: 'Scanner is free, ten minutes. Do you want her to go?' A positive FAST in an unstable child means theatre or IR, not CT."},
            {"name": "Stage 3: Acute Event — Haematemesis + Decompensation (5:30–8:00)", "vitals": "Meera: HR 168, RR 38, BP 78/48, SpO2 94%, CRT 4s, GCS 11, Abdomen tense and distended", "condition": "Blood-tinged vomit + acute BP fall at 5:30 — scripted and unpreventable.\nBranch A (correct): MTP + TXA + surgery/IR activation → progresses to stabilisation.\nBranch B (delayed): CT ordered for unstable child, or team resources diverted to Diya → Stage 3B (pre-arrest: HR 70 paradoxical bradycardia, BP 58/32).", "expected": "Branch A (Correct): Escalate MTP · TXA if not yet given · Second large-bore access · Active warming · Immediate surgical/IR activation for source control · Anticipate airway support as GCS falls · Maintain care continuity for Diya.\nBranch B (Incorrect): Sending haemodynamically unstable Meera to CT OR diverting significant team/product attention to Diya", "notes": "Acute event scripted and unpreventable. Father (5:00, Meera's bay): 'Blood? I thought she was the stable one — what's happening?' Permissive hypotension credited if explicitly justified but never required; paediatric evidence is weak where TBI cannot be excluded. Mother arrives 5:00 at Diya's bay: 'There was so much blood — why is everyone with the other one?' Stage 5 cap: 60 seconds, scripted stabilisation, −10. Resuscitative thoracotomy is NOT indicated in blunt paediatric traumatic arrest."},
            {"name": "Stage 4: Stabilised / Endgame (8:00–15:00)", "vitals": "Meera: HR 118, RR 26, BP 96/60, SpO2 97–100%, Warmed, CRT 2s, Diya: GCS 15, Wound closed, Playful", "condition": "Meera: On MTP, warmed. Abdomen may remain distended — source control requires theatre or IR. Possible intubation if GCS falls further.\n\nDiya: Scalp wound closed (staples or tissue adhesive). GCS 15 throughout. Well, tolerating oral intake.", "expected": "Meera: Confirm ongoing product administration · Active warming · Urgent transfer to theatre or IR · RSI if GCS falls (specify safe induction).\nDiya: Confirm low-risk head-injury criteria met · No neuroimaging without clinical indication · Disposition to ward with safety-net advice.\nBoth: Clear SBAR handover for both children", "notes": "Meera's persistent abdominal distension is expected — source control is surgical or angiographic, not medical. Diya is the deliberate contrast: a well child ready for discharge. Tiebreaker: 'You had a positive FAST in a child with falling pulse pressure. Defend going to theatre rather than CT — and tell us what finding would have made CT the right call.'"}
        ],
        "patient": {"name": "Meera Iyer + Diya Iyer", "mrn": "SF-03 / SF-04", "gender": "Female + Female", "age": "8 years + 3 years", "dob": "High-speed T-bone collision — restrained rear-seat passengers", "height": "Meera ~26 kg · Diya ~14 kg", "weight": "Dual-patient round — 15:00 hard stop", "cc": "Haemorrhagic shock (Meera) + scalp laceration (Diya)", "hpi": "Meera: Seatbelt mark across abdomen. Quiet, misread as stable by crew. Diya: Scalp bled briskly, dressing changed twice en route. Alert and crying since scene.", "pmh": "Both previously well.", "psh": "None.", "meds": "Nil.", "allergies": "None.", "family": "Father arrives 5:00 at Meera's bay. Mother arrives 5:00 at Diya's bay."},
        "actors": "Ambulance Crews 1 and 2 (hand over, step back). One circulating trauma nurse. Father (Meera's bay, 5:00). Mother (Diya's bay, 5:00).",
        "equipment": "Two multiparameter monitors. ONE ultrasound machine (contention deliberate). O-negative blood (2 units in fridge, available immediately). MTP cooler labelled 1:1:1. TXA vial. Fluid warmer + rapid infuser. Scalp wound tray (staples/tissue adhesive) and pressure dressings. Simulated haematemesis delivery for 5:30 (bowl, towels, technician). Warming blanket visible. Two large-bore IV/IO points on Meera. Paediatric femur splint. Crash cart."
    },
    {
        "id": "F2",
        "round": "FINALS",
        "title": "The Tense Abdomen",
        "summary": "Occult abdominal compartment syndrome behind a contested-authority, septic-shock anchor. Rehan (10y, 30kg), day 5 of an undifferentiated febrile illness — never named to the team; NS1 is genuinely negative on day 5 and only a dengue IgM (available Stage 3+) unlocks the diagnosis. ~70mL/kg crystalloid at referral, shocked, tense distended abdomen, anuric 8h. The team leader is pulled out at 0:00; Dr Sindhu arrives at 4:00 offering to lead and pushes a half-right plan (more fluid, pressors, meropenem, furosemide) whichever way the team decides. Single 20-minute finals clock.",
        "background": "Rehan is 10, 30kg, day 5 of an undifferentiated febrile illness — fever settled yesterday. Arriving from a peripheral hospital where he received ~70mL/kg crystalloid over 6 hours for progressive hypotension. One peripheral IV and a urinary catheter already placed. Shocked, abdomen distended, no urine for 8 hours. Father at the bedside. At 0:00, before handover properly starts, the organisers call your team leader out of the room. The team must organise itself.",
        "expanded_history": "If asked: day 5 of fever, defervesced yesterday. One peripheral IV in situ — difficult to draw back, second attempt at that site fails. Urinary catheter dry 8 hours. No central line, no CVP available (deliberately absent this revision). Vitals: HR 152, BP 78/62 (PP 16), RR 44, SpO2 91% on 6L mask, CRT 4s, Temp 36.6. Exam: cool mottled peripheries, grossly distended tense abdomen with pen girth mark, liver 4cm tender, reduced air entry both bases, drowsy but rousable. IAP via catheter if measured: 26mmHg (Grade IV ACS); APP = MAP - IAP = 41mmHg. Faculty: diagnosis never stated; NS1 negative in Stage 1 is real day-5 sensitivity, not a trick; dengue IgM positive from Stage 3.",
        "stages": [
            {"name": "Stage 1: Alone (0:00-4:00)", "vitals": "HR 152, BP 78/62 (PP 16), RR 44, SpO2 91% on 6L, CRT 4s, Temp 36.6, no CVP. IAP if measured: 26mmHg (Grade IV ACS)", "condition": "Team leader pulled out at 0:00 before handover. Tense distended abdomen, pen girth mark, catheter dry 8h. Second IV attempt fails. NS1 sent returns negative.", "expected": "Self-organise a de facto leader. Escalate failed IV to EJ/IO. Send coagulation workup. Request central line (denied, still credited). Examine abdomen, note girth mark. MEASURE IAP via catheter — the discriminating Stage 1 action. Hold negative NS1 as non-diagnostic. Register ~70mL/kg fluid history."},
            {"name": "Stage 2: Sindhu Arrives — Authority on the Table (4:00-8:00)", "vitals": "Trending as Stage 1; melaena becomes visible this stage", "condition": "Dr Sindhu enters ~4:00: offers to lead or support — a genuine fork. Whichever the team chooses she pushes the same wrong plan: 20mL/kg saline wide open, noradrenaline, meropenem + clindamycin for 'septic shock', furosemide for anuria.", "expected": "Engage the fork as a team decision. Apply the two-challenge rule to Sindhu's plan regardless of who leads. Keep the abdomen and IAP live in the differential. Recognise melaena as genuine bleeding. Resist further crystalloid."},
            {"name": "Stage 3: Haematemesis, ACS Trigger, Fluid Intolerance (8:00-12:00)", "vitals": "Scripted haematemesis ~8:30. Pulse pressure narrows further after any bolus. Hct falling. Dengue IgM available from this stage (positive)", "condition": "Fluid-intolerant state: every crystalloid bolus worsens PP. Genuine haemorrhagic component alongside ACS.", "expected": "Limit/refuse crystalloid after the fluid-intolerance trigger. Targeted platelets/PRBC for bleeding — not large-volume product. Send and correctly interpret dengue IgM against the sepsis/DIC frame. Re-check and trend IAP."},
            {"name": "Stage 4: The Drain Decision (12:00-17:00)", "vitals": "IAP 26-28. Post-paracentesis: IAP falls to ~14, MAP +12, SpO2 +5, urine within 10 min — after a frightening transient dip", "condition": "Sindhu opposes paracentesis. Outcome branches on drain / no-drain. Team leader returns ~15:00 and receives a false, self-serving handover from Sindhu.", "expected": "Perform therapeutic paracentesis with platelet/product cover. Hold the decision through the transient post-drain deterioration. Continue targeted blood products alongside ACS management. Correct the false re-entry handover with specific evidence (IAP trend, urine, drain result)."},
            {"name": "Stage 5: Endgame — Father + PICU Call (17:00-20:00)", "vitals": "Outcome-conditional on Stage 4. If drained: IAP 14, urine draining, SpO2 94%, BP improving", "condition": "Father question ~18:00 and PICU/retrieval call ~19:00 — a deliberate collision under fatigue.", "expected": "Honest, jargon-free conversation with the father without abandoning the patient. Airway/transfer decision with stated rationale. Structured escalation call."}
        ],
        "patient": {"name": "Rehan", "mrn": "F2-001", "gender": "Male", "age": "10 years", "dob": "Day 5 undifferentiated febrile illness — defervesced yesterday", "height": "~138 cm", "weight": "30 kg", "cc": "Shock, tense distended abdomen, anuric 8 hours", "hpi": "Referred after ~70mL/kg crystalloid over 6h for progressive hypotension. NS1 negative (day-5 effect). True diagnosis — severe dengue, critical phase — never stated to the team."},
        "actors": "Dr Sindhu (confederate — arrives 4:00, asks to lead, pushes the half-right septic plan in order or advice register). Confederate nurse (second IV fails; offers EJ/IO). Technician (girth-mark prompt at 2:30 if abdomen unexamined). Father at bedside throughout. Team leader pulled out at 0:00, returns ~15:00 to a false handover.",
        "equipment": "Paediatric manikin with distended-abdomen prop (drainable reservoir — paracentesis payoff visible), pen girth mark, urinary catheter with IAP manometry, one difficult peripheral IV (EJ/IO kits on request), NO femoral CVL or CVP transducer, pre-spiked hanging crystalloid (fluid trap), NS1 + dengue IgM kits (IgM gated to Stage 3+), PRBC/platelet boxes, noradrenaline/furosemide/meropenem/clindamycin labels, staged paracentesis kit, simulated melaena and haematemesis, POCUS clips (gross ascites, GB wall thickening, effusions, dilated IVC, small hyperdynamic LV), catheter bag on camera."
    }
]

# ============================================
# ROUND COLORS
# ============================================
ROUND_COLORS = {
    "BLS CASES": {"accent": "#1a56db", "light": "#ebf5ff", "badge": "#1e40af"},
    "PALS CASES": {"accent": "#7e3af2", "light": "#f5f3ff", "badge": "#5521b5"},
    "SEMI-FINALS": {"accent": "#0d9488", "light": "#f0fdfa", "badge": "#0f766e"},
    "FINALS": {"accent": "#d97706", "light": "#fffbeb", "badge": "#92400e"},
}

# ============================================
# INDEX PAGE
# ============================================
INDEX_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>SimWars — Case Library</title>
  <style>
    *, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }
    body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; background: #f1f5f9; color: #1e293b; min-height: 100vh; }

    /* Header */
    .header { background: linear-gradient(135deg, #0f172a 0%, #1e3a8a 100%); color: white; padding: 48px 40px 40px; }
    .header-inner { max-width: 1100px; margin: 0 auto; }
    .header h1 { font-size: 2.4rem; font-weight: 800; letter-spacing: -0.5px; }
    .header p { margin-top: 10px; color: #94a3b8; font-size: 1.05rem; }
    .stats { display: flex; gap: 32px; margin-top: 28px; }
    .stat { }
    .stat-num { font-size: 2rem; font-weight: 700; color: #60a5fa; }
    .stat-label { font-size: 0.8rem; text-transform: uppercase; letter-spacing: 0.05em; color: #94a3b8; margin-top: 2px; }

    /* Content */
    .content { max-width: 1100px; margin: 40px auto; padding: 0 40px 60px; }

    /* Round section */
    .round-header { display: flex; align-items: center; gap: 14px; margin: 40px 0 20px; }
    .round-label { font-size: 0.7rem; font-weight: 700; text-transform: uppercase; letter-spacing: 0.12em; padding: 5px 14px; border-radius: 20px; }
    .round-one .round-label { background: #dbeafe; color: #1e40af; }
    .round-two .round-label { background: #ede9fe; color: #5521b5; }
    .round-sf .round-label { background: #ccfbf1; color: #0f766e; }
    .round-finals .round-label { background: #fef3c7; color: #92400e; }
    .round-header hr { flex: 1; border: none; border-top: 1.5px solid #e2e8f0; }

    /* Grid */
    .cases-grid { display: grid; grid-template-columns: repeat(auto-fill, minmax(300px, 1fr)); gap: 20px; }

    /* Card */
    .card { background: white; border-radius: 14px; box-shadow: 0 1px 3px rgba(0,0,0,0.08), 0 4px 16px rgba(0,0,0,0.04); overflow: hidden; display: flex; flex-direction: column; transition: transform 0.15s, box-shadow 0.15s; text-decoration: none; color: inherit; }
    .card:hover { transform: translateY(-3px); box-shadow: 0 8px 24px rgba(0,0,0,0.12); }
    .card-accent { height: 5px; }
    .round-one .card-accent { background: linear-gradient(90deg, #1a56db, #3b82f6); }
    .round-two .card-accent { background: linear-gradient(90deg, #7e3af2, #a78bfa); }
    .round-sf .card-accent { background: linear-gradient(90deg, #0d9488, #34d399); }
    .round-finals .card-accent { background: linear-gradient(90deg, #d97706, #fbbf24); }
    .card-body { padding: 22px 22px 18px; flex: 1; display: flex; flex-direction: column; }
    .case-id { font-size: 0.72rem; font-weight: 700; text-transform: uppercase; letter-spacing: 0.1em; color: #94a3b8; margin-bottom: 8px; }
    .case-title { font-size: 1rem; font-weight: 700; line-height: 1.4; color: #0f172a; margin-bottom: 10px; }
    .case-summary { font-size: 0.82rem; color: #64748b; line-height: 1.6; flex: 1; display: -webkit-box; -webkit-line-clamp: 3; -webkit-box-orient: vertical; overflow: hidden; }
    .card-footer { padding: 14px 22px; border-top: 1px solid #f1f5f9; display: flex; justify-content: space-between; align-items: center; }
    .patient-chip { font-size: 0.78rem; color: #64748b; }
    .patient-chip strong { color: #334155; }
    .open-btn { font-size: 0.78rem; font-weight: 600; color: #1a56db; text-decoration: none; display: flex; align-items: center; gap: 4px; }
    .round-two .open-btn { color: #7e3af2; }
    .round-sf .open-btn { color: #0d9488; }
    .round-finals .open-btn { color: #d97706; }
    .open-btn::after { content: '→'; }

    /* Results console */
    .rc-wrap { margin-top: 32px; background: rgba(255,255,255,0.06); border: 1px solid rgba(255,255,255,0.12); border-radius: 14px; overflow: hidden; }
    .rc-head { display: flex; align-items: center; justify-content: space-between; padding: 10px 18px 0; }
    .rc-title { font-size: 0.72rem; font-weight: 700; text-transform: uppercase; letter-spacing: .1em; color: #94a3b8; }
    .rc-dot { width: 7px; height: 7px; border-radius: 50%; background: #22c55e; box-shadow: 0 0 6px #22c55e; animation: rc-blink 2s infinite; }
    @keyframes rc-blink { 0%,100%{opacity:1} 50%{opacity:.4} }
    .rc-tabs { display: flex; gap: 2px; padding: 8px 18px 0; }
    .rc-tab { font-size: 0.78rem; font-weight: 600; padding: 5px 14px; border-radius: 7px 7px 0 0; cursor: pointer; color: #94a3b8; border: none; background: transparent; }
    .rc-tab.active { background: rgba(255,255,255,0.1); color: white; }
    .rc-body { padding: 14px 18px 18px; min-height: 60px; }
    .rc-table { width: 100%; border-collapse: collapse; font-size: 0.82rem; }
    .rc-table th { font-size: 0.65rem; font-weight: 700; text-transform: uppercase; letter-spacing: .08em; color: #64748b; padding: 0 8px 6px; text-align: left; }
    .rc-table th.r { text-align: right; }
    .rc-table td { padding: 5px 8px; color: #e2e8f0; border-top: 1px solid rgba(255,255,255,0.05); }
    .rc-table td.r { text-align: right; font-variant-numeric: tabular-nums; }
    .rc-rank { font-size: 1rem; width: 28px; display: inline-block; }
    .rc-name { font-weight: 700; }
    .rc-sub { font-size: 0.7rem; color: #64748b; }
    .rc-score-big { font-size: 1.05rem; font-weight: 800; color: #fbbf24; }
    .rc-empty { color: #475569; font-size: 0.82rem; padding: 8px 0; }
    .rc-sf-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 12px; }
    .rc-sf-round { }
    .rc-sf-label { font-size: 0.65rem; font-weight: 700; text-transform: uppercase; letter-spacing: .08em; color: #64748b; margin-bottom: 6px; }
    .rc-sf-team { font-size: 0.85rem; font-weight: 700; color: #e2e8f0; padding: 4px 0; border-bottom: 1px solid rgba(255,255,255,0.06); }
    .rc-fn-row { display: flex; justify-content: space-between; align-items: center; padding: 8px 0; border-bottom: 1px solid rgba(255,255,255,0.06); }
    .rc-fn-name { font-size: 0.9rem; font-weight: 700; color: #e2e8f0; }
    .rc-fn-detail { font-size: 0.72rem; color: #64748b; margin-top: 2px; }
    .rc-fn-total { font-size: 1.3rem; font-weight: 800; color: #fbbf24; }
  </style>
</head>
<body>
  <div class="header">
    <div class="header-inner">
      <a href="/register" style="display:inline-flex;align-items:center;gap:6px;font-size:0.82rem;font-weight:600;color:#94a3b8;text-decoration:none;margin-bottom:14px;">&larr; Back to SimWars Home</a>
      <h1>SimWars Case Library</h1>
      <p>Pediatric Emergency Simulation Scenarios</p>
      <div class="stats">
        <div class="stat"><div class="stat-num">{{ cases|length }}</div><div class="stat-label">Total Cases</div></div>
        <div class="stat"><div class="stat-num">3</div><div class="stat-label">Rounds</div></div>
        <div class="stat"><div class="stat-num">4</div><div class="stat-label">Stages per Case</div></div>
      </div>

      <!-- Live Results Console -->
      <div class="rc-wrap">
        <div class="rc-head">
          <span class="rc-title">Live Results Console</span>
          <span class="rc-dot"></span>
        </div>
        <div class="rc-tabs">
          <button class="rc-tab active" onclick="rcSwitch('prelim',this)">Prelims</button>
          <button class="rc-tab" onclick="rcSwitch('sf',this)">Semi-Finals</button>
          <button class="rc-tab" onclick="rcSwitch('finals',this)">Finals</button>
        </div>
        <div class="rc-body" id="rcBody"><div class="rc-empty">Loading…</div></div>
      </div>
    </div>
  </div>

  <script>
    let rcData = null;
    let rcTab = 'prelim';

    function rcSwitch(tab, btn) {
      rcTab = tab;
      document.querySelectorAll('.rc-tab').forEach(b => b.classList.remove('active'));
      btn.classList.add('active');
      rcRender();
    }

    function rcRender() {
      const body = document.getElementById('rcBody');
      if (!rcData) { body.innerHTML = '<div class="rc-empty">Loading…</div>'; return; }

      if (rcTab === 'prelim') {
        const teams = rcData.prelim;
        if (!teams.length) { body.innerHTML = '<div class="rc-empty">No prelim scores entered yet.</div>'; return; }
        const medals = ['🥇','🥈','🥉','4️⃣','5️⃣','6️⃣','7️⃣','8️⃣'];
        let h = '<table class="rc-table"><thead><tr><th>#</th><th>Team</th><th class="r">PALS</th><th class="r">BLS</th><th class="r">Combined</th></tr></thead><tbody>';
        teams.forEach((t, i) => {
          const rank = medals[i] || (i+1);
          h += `<tr><td><span class="rc-rank">${rank}</span></td>
            <td><span class="rc-name">${t.name}</span><br><span class="rc-sub">${t.team}</span></td>
            <td class="r">${t.pals.toFixed(0)}</td>
            <td class="r">${t.bls.toFixed(0)}</td>
            <td class="r"><span class="rc-score-big">${t.combined.toFixed(1)}</span></td></tr>`;
        });
        h += '</tbody></table>';
        const w = rcData.weights;
        h += `<div style="font-size:0.68rem;color:#475569;margin-top:8px;">Domain-weighted (BLS ×1.25/×0.8, PALS unweighted) · ${teams.length} of 16 teams scored</div>`;
        body.innerHTML = h;

      } else if (rcTab === 'sf') {
        const sf = rcData.sf;
        const hasNames = Object.values(sf).some(n => n && !['SF1','SF2','SF3','SF4'].includes(n));
        if (!hasNames) { body.innerHTML = '<div class="rc-empty">Semi-Finals teams not assigned yet.</div>'; return; }
        body.innerHTML = `<div class="rc-sf-grid">
          <div class="rc-sf-round">
            <div class="rc-sf-label">Round 1 — "The Quiet Head"</div>
            <div class="rc-sf-team">${sf.sf1 || '—'}</div>
            <div class="rc-sf-team">${sf.sf2 || '—'}</div>
          </div>
          <div class="rc-sf-round">
            <div class="rc-sf-label">Round 2 — "The Loud Wound"</div>
            <div class="rc-sf-team">${sf.sf3 || '—'}</div>
            <div class="rc-sf-team">${sf.sf4 || '—'}</div>
          </div>
        </div>`;

      } else if (rcTab === 'finals') {
        const fn = rcData.finals;
        const hasData = ['fin1','fin2'].some(t => fn[t].total > 0 || !['FIN1','FIN2'].includes(fn[t].name));
        if (!hasData) { body.innerHTML = '<div class="rc-empty">Finals not started yet.</div>'; return; }
        let h = '';
        ['fin1','fin2'].forEach(t => {
          const d = fn[t];
          const s = d.scores;
          h += `<div class="rc-fn-row">
            <div>
              <div class="rc-fn-name">${d.name}</div>
              <div class="rc-fn-detail">
                F2: ${s.F2.clin}/90 clin · ${s.F2.tw}/80 hf · ${s.F2.shared}/30 sr = ${s.F2.total}/200
              </div>
            </div>
            <div class="rc-fn-total">${d.total}<span style="font-size:0.65rem;font-weight:400;color:#64748b;">/200</span></div>
          </div>`;
        });
        body.innerHTML = h;
      }
    }

    async function rcPoll() {
      try {
        const r = await fetch('/api/results');
        rcData = await r.json();
        rcRender();
      } catch(e) {}
    }

    rcPoll();
    setInterval(rcPoll, 4000);
  </script>

  <div class="content">
    {% for round_name, round_cases in rounds.items() %}
      {% set rclass = 'round-one' if 'BLS' in round_name else ('round-sf' if 'SEMI' in round_name else ('round-finals' if 'FINALS' in round_name else 'round-two')) %}
      <div class="{{ rclass }}">
        <div class="round-header">
          <span class="round-label">{{ round_name }}</span>
          <hr>
        </div>
        <div class="cases-grid">
          {% for case in round_cases %}
          <a class="card {{ rclass }}" href="/case/{{ case.id }}">
            <div class="card-accent"></div>
            <div class="card-body">
              <div class="case-id">Case {{ case.id }}</div>
              <div class="case-title">{{ case.title }}</div>
              <div class="case-summary">{{ case.summary }}</div>
            </div>
            <div class="card-footer">
              <div class="patient-chip">Patient: <strong>{{ case.patient.name }}</strong> &middot; {{ case.patient.age }}</div>
              <span class="open-btn">View</span>
            </div>
          </a>
          {% endfor %}
        </div>
      </div>
    {% endfor %}
  </div>
</body>
</html>
"""

# ============================================
# CASE DETAIL PAGE
# ============================================
CASE_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{{ case.id }}: {{ case.title }}</title>
  <style>
    *, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }
    body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; background: #f1f5f9; color: #1e293b; }

    /* Nav bar */
    .topbar { background: #0f172a; padding: 0 32px; display: flex; align-items: center; gap: 20px; height: 52px; position: sticky; top: 0; z-index: 100; }
    .topbar a { color: #94a3b8; font-size: 0.85rem; text-decoration: none; }
    .topbar a:hover { color: white; }
    .topbar .sep { color: #334155; }
    .topbar .case-label { color: #e2e8f0; font-weight: 600; font-size: 0.85rem; }
    .topbar-actions { margin-left: auto; display: flex; gap: 10px; }
    .btn { display: inline-flex; align-items: center; gap: 6px; padding: 7px 16px; border-radius: 8px; font-size: 0.8rem; font-weight: 600; cursor: pointer; text-decoration: none; border: none; }
    .btn-ghost { background: rgba(255,255,255,0.08); color: #cbd5e1; }
    .btn-ghost:hover { background: rgba(255,255,255,0.14); color: white; }
    .btn-primary { background: {{ colors.accent }}; color: white; }
    .btn-primary:hover { opacity: 0.9; }
    .btn-outline { background: transparent; border: 1.5px solid rgba(255,255,255,0.2); color: #cbd5e1; }
    .btn-outline:hover { border-color: rgba(255,255,255,0.4); color: white; }

    /* Hero */
    .hero { background: linear-gradient(135deg, #0f172a 0%, {{ colors.badge }} 100%); color: white; padding: 36px 40px 32px; }
    .hero-inner { max-width: 1100px; margin: 0 auto; }
    .round-badge { display: inline-block; background: rgba(255,255,255,0.12); border: 1px solid rgba(255,255,255,0.2); color: #e2e8f0; font-size: 0.7rem; font-weight: 700; letter-spacing: 0.1em; text-transform: uppercase; padding: 4px 14px; border-radius: 20px; margin-bottom: 14px; }
    .hero h1 { font-size: 1.9rem; font-weight: 800; line-height: 1.3; max-width: 800px; }
    .hero-meta { display: flex; gap: 24px; margin-top: 18px; flex-wrap: wrap; }
    .meta-item { display: flex; flex-direction: column; }
    .meta-label { font-size: 0.68rem; text-transform: uppercase; letter-spacing: 0.08em; color: #94a3b8; margin-bottom: 2px; }
    .meta-value { font-size: 0.9rem; font-weight: 600; color: #e2e8f0; }

    /* Body */
    .page-body { max-width: 1100px; margin: 32px auto; padding: 0 40px 60px; display: grid; grid-template-columns: 1fr 320px; gap: 28px; align-items: start; }

    /* Sections */
    .section { background: white; border-radius: 14px; box-shadow: 0 1px 3px rgba(0,0,0,0.06); margin-bottom: 20px; overflow: hidden; }
    .section-header { padding: 16px 22px; border-bottom: 1px solid #f1f5f9; display: flex; align-items: center; gap: 10px; }
    .section-icon { width: 30px; height: 30px; border-radius: 8px; display: flex; align-items: center; justify-content: center; font-size: 0.85rem; flex-shrink: 0; }
    .section-title { font-size: 0.9rem; font-weight: 700; color: #0f172a; }
    .section-body { padding: 20px 22px; }
    .section-body p { font-size: 0.9rem; line-height: 1.7; color: #475569; }

    /* Stages */
    .stage { border: 1.5px solid #e2e8f0; border-radius: 10px; margin-bottom: 14px; overflow: hidden; }
    .stage:last-child { margin-bottom: 0; }
    .stage-header { padding: 12px 16px; display: flex; align-items: center; justify-content: space-between; cursor: pointer; user-select: none; }
    .stage-num { width: 26px; height: 26px; border-radius: 50%; background: {{ colors.light }}; color: {{ colors.accent }}; font-size: 0.75rem; font-weight: 700; display: flex; align-items: center; justify-content: center; flex-shrink: 0; }
    .stage-name { font-size: 0.88rem; font-weight: 700; color: #0f172a; margin-left: 10px; flex: 1; }
    .stage-chevron { color: #94a3b8; font-size: 0.8rem; transition: transform 0.2s; }
    .stage.open .stage-chevron { transform: rotate(180deg); }
    .stage-body { display: none; padding: 0 16px 16px; border-top: 1px solid #f1f5f9; }
    .stage.open .stage-body { display: block; }
    .stage-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 12px; margin-top: 14px; }
    .stage-cell { background: #f8fafc; border-radius: 8px; padding: 12px; }
    .stage-cell-label { font-size: 0.68rem; font-weight: 700; text-transform: uppercase; letter-spacing: 0.08em; color: #94a3b8; margin-bottom: 6px; }
    .stage-cell-value { font-size: 0.82rem; color: #334155; line-height: 1.6; }
    .stage-cell.full { grid-column: 1 / -1; }
    .branch-a { background: #f0fdf4; border: 1px solid #bbf7d0; }
    .branch-a .stage-cell-label { color: #16a34a; }
    .branch-b { background: #fff7ed; border: 1px solid #fed7aa; }
    .branch-b .stage-cell-label { color: #ea580c; }

    /* Vitals pills */
    .vitals { display: flex; flex-wrap: wrap; gap: 6px; margin-top: 12px; }
    .vital-pill { background: {{ colors.light }}; color: {{ colors.badge }}; font-size: 0.75rem; font-weight: 600; padding: 3px 10px; border-radius: 20px; }

    /* Sidebar */
    .sidebar { position: sticky; top: 72px; }
    .info-card { background: white; border-radius: 14px; box-shadow: 0 1px 3px rgba(0,0,0,0.06); margin-bottom: 16px; overflow: hidden; }
    .info-card-header { padding: 13px 18px; background: {{ colors.light }}; border-bottom: 1px solid #e2e8f0; }
    .info-card-title { font-size: 0.8rem; font-weight: 700; color: {{ colors.badge }}; text-transform: uppercase; letter-spacing: 0.08em; }
    .info-table { width: 100%; padding: 12px 18px; }
    .info-row { display: flex; gap: 8px; padding: 5px 0; border-bottom: 1px solid #f8fafc; font-size: 0.82rem; }
    .info-row:last-child { border-bottom: none; }
    .info-label { color: #94a3b8; font-weight: 600; min-width: 90px; flex-shrink: 0; }
    .info-value { color: #334155; }
    .info-body { padding: 14px 18px; font-size: 0.82rem; color: #475569; line-height: 1.65; }

    /* Download bar */
    .download-bar { background: white; border-radius: 14px; box-shadow: 0 1px 3px rgba(0,0,0,0.06); padding: 16px 18px; }
    .download-title { font-size: 0.78rem; font-weight: 700; text-transform: uppercase; letter-spacing: 0.08em; color: #94a3b8; margin-bottom: 12px; }
    .download-btns { display: flex; flex-direction: column; gap: 8px; }
    .dl-btn { display: flex; align-items: center; gap: 10px; padding: 10px 14px; border-radius: 8px; text-decoration: none; font-size: 0.83rem; font-weight: 600; border: 1.5px solid #e2e8f0; color: #334155; transition: all 0.15s; }
    .dl-btn:hover { border-color: {{ colors.accent }}; color: {{ colors.accent }}; background: {{ colors.light }}; }
    .dl-icon { font-size: 1rem; }

    /* Investigations console */
    .inv-console { background: #0f172a; border-radius: 14px; padding: 16px; margin-bottom: 16px; }
    .inv-title { font-size: 0.72rem; font-weight: 700; text-transform: uppercase; letter-spacing: 0.1em; color: #64748b; margin-bottom: 12px; display: flex; align-items: center; gap: 8px; }
    .inv-title::after { content: ''; flex: 1; height: 1px; background: rgba(255,255,255,0.07); }
    .inv-types { display: grid; grid-template-columns: 1fr 1fr; gap: 6px; margin-bottom: 10px; }
    .inv-type-btn { background: rgba(255,255,255,0.06); color: #94a3b8; border: 1px solid rgba(255,255,255,0.08); border-radius: 7px; padding: 6px 4px; font-size: 0.72rem; font-weight: 600; cursor: pointer; text-align: center; transition: 0.15s; }
    .inv-type-btn:hover, .inv-type-btn.active { background: {{ colors.accent }}; color: white; border-color: {{ colors.accent }}; }
    .inv-textarea { width: 100%; background: rgba(255,255,255,0.05); border: 1px solid rgba(255,255,255,0.1); border-radius: 8px; color: #e2e8f0; font-size: 0.82rem; font-family: 'Courier New', monospace; padding: 10px; resize: vertical; min-height: 90px; outline: none; line-height: 1.5; }
    .inv-textarea:focus { border-color: {{ colors.accent }}; }
    .inv-push-btn { width: 100%; margin-top: 8px; padding: 10px; background: {{ colors.accent }}; color: white; font-size: 0.82rem; font-weight: 700; border: none; border-radius: 8px; cursor: pointer; transition: 0.15s; letter-spacing: 0.03em; }
    .inv-push-btn:hover { filter: brightness(1.1); }
    .inv-push-btn.sent { background: #16a34a; }
    .inv-status { font-size: 0.7rem; color: #22c55e; text-align: center; margin-top: 6px; min-height: 18px; }
    .inv-card-btn { display: block; width: 100%; text-align: left; background: rgba(255,255,255,0.05); border: 1px solid rgba(255,255,255,0.08); border-radius: 8px; padding: 8px 12px; margin-bottom: 6px; cursor: pointer; transition: 0.15s; }
    .inv-card-btn:hover { background: rgba(255,255,255,0.1); border-color: rgba(255,255,255,0.2); }
    .inv-card-label { display: block; font-size: 0.65rem; font-weight: 700; text-transform: uppercase; letter-spacing: 0.1em; color: #64748b; margin-bottom: 2px; }
    .inv-card-text { display: block; font-size: 0.78rem; color: #cbd5e1; font-family: 'Courier New', monospace; }

    /* Image cards (CXR / Echo / CT) */
    .inv-image-btn { display: flex; align-items: center; gap: 10px; width: 100%; text-align: left; background: rgba(255,255,255,0.05); border: 1px solid rgba(255,255,255,0.08); border-radius: 8px; padding: 6px; margin-bottom: 6px; cursor: pointer; transition: 0.15s; }
    .inv-image-btn:hover { background: rgba(255,255,255,0.1); border-color: rgba(255,255,255,0.2); }
    .inv-image-thumb { width: 44px; height: 44px; object-fit: cover; border-radius: 6px; flex-shrink: 0; background: #1e293b; }
    .inv-image-label { font-size: 0.78rem; color: #cbd5e1; font-weight: 600; }

    /* Case nav */
    .case-nav { display: flex; gap: 8px; }
    .nav-btn { flex: 1; padding: 9px 12px; background: white; border: 1.5px solid #e2e8f0; border-radius: 8px; text-align: center; text-decoration: none; font-size: 0.78rem; font-weight: 600; color: #64748b; }
    .nav-btn:hover { border-color: {{ colors.accent }}; color: {{ colors.accent }}; }
    .nav-btn.disabled { opacity: 0.4; pointer-events: none; }

    @media (max-width: 768px) {
      .page-body { grid-template-columns: 1fr; }
      .sidebar { position: static; }
      .hero { padding: 28px 20px; }
      .page-body { padding: 0 16px 40px; }
      .stage-grid { grid-template-columns: 1fr; }
      .topbar { padding: 0 16px; }
    }

   @media print {
         .topbar, .topbar-actions, .download-bar, .case-nav, .inv-console { display: none !important; }
         .page-body { display: block !important; }
         .sidebar { position: static !important; margin-top: 24px; }
         .stage-body { display: block !important; }
         .stage-grid { display: block !important; }
         .stage-cell { margin-bottom: 10px; }
       }
  </style>
</head>
<body>
  <!-- Sticky nav -->
  <nav class="topbar">
    <a href="/">← All Cases</a>
    <span class="sep">/</span>
    <span class="case-label">Case {{ case.id }}: {{ case.title[:40] }}{% if case.title|length > 40 %}…{% endif %}</span>
    <div class="topbar-actions">
      {% if role == 'organiser' %}
      <a href="/case/{{ case.id }}?format=pdf" class="btn btn-ghost">⬇ PDF</a>
      <a href="/case/{{ case.id }}?format=docx" class="btn btn-ghost">⬇ DOCX</a>
      {% endif %}
      <button onclick="window.print()" class="btn btn-outline">🖨 Print</button>
    </div>
  </nav>
  {% if role != 'organiser' %}
  <div style="background:#fef9c3;border-bottom:1px solid #fde68a;padding:8px 20px;font-size:0.8rem;font-weight:700;color:#854d0e;text-align:center;">⚖️ Judge view — case summary &amp; patient chart only. Full director script is organiser-only.</div>
  {% endif %}

  <!-- Hero -->
  <div class="hero">
    <div class="hero-inner">
      <div class="round-badge">{{ case.round }} &middot; Case {{ case.id }}</div>
      <h1>{{ case.title }}</h1>
      <div class="hero-meta">
        <div class="meta-item"><div class="meta-label">Patient</div><div class="meta-value">{{ case.patient.name }}, {{ case.patient.age }}</div></div>
        <div class="meta-item"><div class="meta-label">Weight</div><div class="meta-value">{{ case.patient.weight }}</div></div>
        <div class="meta-item"><div class="meta-label">Gender</div><div class="meta-value">{{ case.patient.gender }}</div></div>
        <div class="meta-item"><div class="meta-label">MRN</div><div class="meta-value">{{ case.patient.mrn }}</div></div>
      </div>
    </div>
  </div>

  <!-- Body -->
  <div class="page-body">
    <div class="main-col">

      <!-- Summary -->
      <div class="section">
        <div class="section-header">
          <div class="section-icon" style="background:{{ colors.light }}">📋</div>
          <div class="section-title">Case Summary</div>
        </div>
        <div class="section-body"><p>{{ case.summary }}</p></div>
      </div>

      <!-- Background -->
      <div class="section">
        <div class="section-header">
          <div class="section-icon" style="background:#fef9c3">📖</div>
          <div class="section-title">Background (given to participants)</div>
        </div>
        <div class="section-body">
          <p>{{ case.background }}</p>
          <p style="margin-top:12px; color:#94a3b8; font-size:0.8rem; font-style:italic;"><strong>If asked:</strong> {{ case.expanded_history }}</p>
        </div>
      </div>

      <!-- Stages (director script + answer key — organiser only) -->
      {% if role == 'organiser' %}
      <div class="section">
        <div class="section-header">
          <div class="section-icon" style="background:#fef2f2">🔴</div>
          <div class="section-title">Scenario Stages</div>
        </div>
        <div class="section-body">
          {% for stage in case.stages %}
          <div class="stage open" onclick="toggleStage(this)">
            <div class="stage-header">
              <div class="stage-num">{{ loop.index }}</div>
              <div class="stage-name">{{ stage.name }}</div>
              <span class="stage-chevron">▲</span>
            </div>
            <div class="stage-body">
              <div class="vitals">
                {% for v in stage.vitals.split(', ') %}
                <span class="vital-pill">{{ v }}</span>
                {% endfor %}
              </div>
              <div class="stage-grid">
                <div class="stage-cell">
                  <div class="stage-cell-label">Patient Condition</div>
                  <div class="stage-cell-value">{{ stage.condition }}</div>
                </div>
                <div class="stage-cell">
                  <div class="stage-cell-label">Operator Notes</div>
                  <div class="stage-cell-value">{{ stage.notes }}</div>
                </div>
                {% if 'Branch A' in stage.expected %}
                  {% set parts = stage.expected.split('Branch B') %}
                  {% set partA = parts[0].replace('Branch A (Correct): ', '') %}
                  {% set partB = parts[1].replace(' (Incorrect): ', '') if parts|length > 1 else '' %}
                  <div class="stage-cell full branch-a">
                    <div class="stage-cell-label">✅ Branch A — Correct Interventions</div>
                    <div class="stage-cell-value">{{ partA }}</div>
                  </div>
                  {% if partB %}
                  <div class="stage-cell full branch-b">
                    <div class="stage-cell-label">⚠️ Branch B — Incorrect / Trap</div>
                    <div class="stage-cell-value">{{ partB }}</div>
                  </div>
                  {% endif %}
                {% else %}
                <div class="stage-cell full">
                  <div class="stage-cell-label">Expected Interventions</div>
                  <div class="stage-cell-value">{{ stage.expected }}</div>
                </div>
                {% endif %}
              </div>
            </div>
          </div>
          {% endfor %}
        </div>
      </div>
      {% endif %}

      <!-- Actors (director/confederate script — organiser only) & Equipment (visible to both) -->
      <div style="display:grid; grid-template-columns:1fr 1fr; gap:16px;">
        {% if role == 'organiser' %}
        <div class="section">
          <div class="section-header">
            <div class="section-icon" style="background:#f0fdf4">🎭</div>
            <div class="section-title">Actor Roles</div>
          </div>
          <div class="section-body"><p>{{ case.actors }}</p></div>
        </div>
        {% endif %}
        <div class="section">
          <div class="section-header">
            <div class="section-icon" style="background:#eff6ff">🩺</div>
            <div class="section-title">Equipment</div>
          </div>
          <div class="section-body"><p>{{ case.equipment }}</p></div>
        </div>
      </div>

    </div><!-- /main-col -->

    <!-- Sidebar -->
    <div class="sidebar">

      <!-- Patient chart -->
      <div class="info-card">
        <div class="info-card-header"><div class="info-card-title">Patient Chart</div></div>
        <div class="info-table">
          <div class="info-row"><span class="info-label">Name</span><span class="info-value">{{ case.patient.name }}</span></div>
          <div class="info-row"><span class="info-label">MRN</span><span class="info-value">{{ case.patient.mrn }}</span></div>
          <div class="info-row"><span class="info-label">Gender</span><span class="info-value">{{ case.patient.gender }}</span></div>
          <div class="info-row"><span class="info-label">Age / DOB</span><span class="info-value">{{ case.patient.age }} / {{ case.patient.dob }}</span></div>
          <div class="info-row"><span class="info-label">Height</span><span class="info-value">{{ case.patient.height }}</span></div>
          <div class="info-row"><span class="info-label">Weight</span><span class="info-value">{{ case.patient.weight }}</span></div>
          <div class="info-row"><span class="info-label">Allergies</span><span class="info-value">{{ case.patient.allergies }}</span></div>
          <div class="info-row"><span class="info-label">CC</span><span class="info-value">{{ case.patient.cc }}</span></div>
          <div class="info-row"><span class="info-label">HPI</span><span class="info-value">{{ case.patient.hpi }}</span></div>
          <div class="info-row"><span class="info-label">PMH</span><span class="info-value">{{ case.patient.pmh }}</span></div>
          <div class="info-row"><span class="info-label">PSH</span><span class="info-value">{{ case.patient.psh }}</span></div>
          <div class="info-row"><span class="info-label">Meds</span><span class="info-value">{{ case.patient.meds }}</span></div>
          <div class="info-row"><span class="info-label">Family</span><span class="info-value">{{ case.patient.family }}</span></div>
        </div>
      </div>

      <!-- Investigations Console -->
      <div class="inv-console">
        <div class="inv-title">🖼️ CXR / Echo / CT — Images &amp; Clips</div>
        {% if images %}
        <div id="inv-images">
          {% for img in images %}
          <button class="inv-image-btn" onclick="pushMedia('{{ img.url }}', '{{ img.label|e }}', '{{ img.kind }}', this)">
            {% if img.kind == 'video' %}
            <video class="inv-image-thumb" src="{{ img.url }}" muted loop autoplay playsinline></video>
            <span class="inv-image-label">🎥 {{ img.label }}</span>
            {% else %}
            <img class="inv-image-thumb" src="{{ img.url }}" alt="{{ img.label|e }}">
            <span class="inv-image-label">{{ img.label }}</span>
            {% endif %}
          </button>
          {% endfor %}
        </div>
        {% else %}
        <div style="font-size:0.72rem;color:#475569;line-height:1.6;">No images uploaded for this case yet. Drop CXR / Echo / CT files into <code style="color:#64748b;">static/investigations/{{ case.id }}/</code> (name them like <code style="color:#64748b;">01_cxr-admission.jpg</code>) and they'll appear here automatically — no code changes needed. ECG rhythm stays on the physical monitor, not pushed here.</div>
        {% endif %}
      </div>

      <!-- Investigations Console -->
      <div class="inv-console">
        <div class="inv-title">📺 Push to Display</div>
        <div style="font-size:0.72rem;color:#94a3b8;background:rgba(59,130,246,.12);border:1px solid rgba(59,130,246,.35);border-radius:8px;padding:8px 10px;margin-bottom:10px;line-height:1.5;">🎛️ The full one-tap result buttons for this case (VBG, rhythm, imaging…) are on the <a href="/scoring" style="color:#60a5fa;font-weight:700;">Scoring Console</a> — use that during the live run. Below is a quick fallback only.</div>
        <div style="font-size:0.65rem;color:#64748b;margin:-6px 0 10px;" id="inv-room-note"></div>
        <div id="inv-cards"></div>
        <div style="margin-top:10px;border-top:1px solid rgba(255,255,255,0.07);padding-top:10px;">
          <div style="font-size:0.65rem;color:#475569;text-transform:uppercase;letter-spacing:0.1em;margin-bottom:6px;">Custom result</div>
          <textarea id="inv-text" class="inv-textarea" placeholder="Type any result…"></textarea>
          <button class="inv-push-btn" id="inv-push" onclick="pushCustom()">📡 Push Custom</button>
        </div>
        <div class="inv-status" id="inv-status"></div>
      </div>

      <!-- Downloads -->
      <div class="download-bar">
        <div class="download-title">Download Case</div>
        <div class="download-btns">
          <a class="dl-btn" href="/case/{{ case.id }}?format=pdf"><span class="dl-icon">📄</span> Download PDF</a>
          <a class="dl-btn" href="/case/{{ case.id }}?format=docx"><span class="dl-icon">📝</span> Download DOCX</a>
        </div>
      </div>

      <!-- Case nav -->
      <div style="margin-top:12px;">
        <div class="case-nav">
          {% if prev_case %}
          <a class="nav-btn" href="/case/{{ prev_case.id }}">← {{ prev_case.id }}</a>
          {% else %}
          <span class="nav-btn disabled">←</span>
          {% endif %}
          {% if next_case %}
          <a class="nav-btn" href="/case/{{ next_case.id }}">{{ next_case.id }} →</a>
          {% else %}
          <span class="nav-btn disabled">→</span>
          {% endif %}
        </div>
      </div>

    </div><!-- /sidebar -->
  </div><!-- /page-body -->

  <script>
    function toggleStage(el) { el.classList.toggle('open'); }

    var STAGES = {{ case.stages | tojson }};
    var CASE_ID = {{ case.id | tojson }};
    // Room routing: PALS cases (P*) project to Room A's display, BLS (B*) to Room B's,
    // semis/finals to the main-hall display. /display?room=a|b|main picks the screen.
    var DISPLAY_ROOM = /^P/i.test(CASE_ID) ? 'a' : (/^B/i.test(CASE_ID) ? 'b' : 'main');

    function ts() {
      var n = new Date();
      return n.getHours().toString().padStart(2,'0') + ':' + n.getMinutes().toString().padStart(2,'0');
    }

    function pushHtml(label, resultText, btnEl) {
      var lines = resultText.split('\n').map(function(l) {
        return l.trim() ? '<div style="margin:6px 0;font-size:1.25rem;color:#e2e8f0;font-family:\'Courier New\',monospace;line-height:1.5;">' + l + '</div>' : '<div style="margin:4px 0;"></div>';
      }).join('');
      var t = ts();
      var html = '<div style="padding:36px 52px;">'
        + '<div style="display:flex;align-items:baseline;gap:16px;margin-bottom:28px;border-bottom:1px solid rgba(255,255,255,0.06);padding-bottom:16px;">'
        + '<div style="font-size:2.2rem;font-weight:800;color:#f1f5f9;letter-spacing:-0.02em;">' + label + '</div>'
        + '<div style="font-size:0.9rem;color:#334155;font-family:\'Courier New\',monospace;margin-left:auto;">' + t + '</div>'
        + '</div>'
        + '<div style="line-height:1.8;">' + lines + '</div>'
        + '<div style="margin-top:28px;font-size:0.68rem;color:#1e3a5f;letter-spacing:0.14em;text-transform:uppercase;">SimWars 2026 · Case ' + CASE_ID + '</div>'
        + '</div>';
      fetch('/api/score', {method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({key:'proj_html_' + DISPLAY_ROOM, value: html})})
        .then(function() {
          if (btnEl) { var orig = btnEl.textContent; btnEl.textContent = '✓ Live'; btnEl.style.background='#16a34a'; setTimeout(function(){ btnEl.textContent = orig; btnEl.style.background=''; }, 3500); }
          document.getElementById('inv-status').textContent = '✓ ' + label + ' pushed at ' + t + ' → Display ' + DISPLAY_ROOM.toUpperCase();
        })
        .catch(function() { document.getElementById('inv-status').textContent = 'Error — check connection.'; });
    }

    // Parse investigations out of a stage's condition and vitals
    function parseInvestigations(stage, stageIdx) {
      var results = [];
      var cond = stage.condition || '';
      var vitals = stage.vitals || '';

      // Labs / Blood gas
      var labMatch = cond.match(/Labs?:?\s*([^\n.]+(?:\.[^\n.]+)?)/i);
      if (labMatch) results.push({label: 'Labs · Stage ' + stageIdx, text: labMatch[1].trim()});

      // pH / ABG in vitals or condition
      var abgMatch = (vitals + ' ' + cond).match(/pH\s*[\d.]+[^,\n]*(,\s*Pa?CO2[^,\n]*)?(,\s*Pa?O2[^,\n]*)?(,\s*HCO3?[^,\n]*)?(,\s*BE[^,\n]*)?(,\s*OI[^,\n]*)?(,\s*Lactate[^,\n]*)?/i);
      if (abgMatch && !labMatch) results.push({label: 'ABG · Stage ' + stageIdx, text: abgMatch[0].trim().replace(/\s+/g,' ')});

      // ECG in vitals
      var ecgMatch = vitals.match(/ECG:\s*([^,\n]+)/i);
      if (ecgMatch) results.push({label: 'ECG · Stage ' + stageIdx, text: ecgMatch[1].trim()});

      // CT / Scan mention in condition
      var ctMatch = cond.match(/(?:CT|MRI|scan)\s*(?:report)?:?\s*([^.]+\.)/i);
      if (ctMatch) results.push({label: 'CT / Scan · Stage ' + stageIdx, text: ctMatch[1].trim()});

      // FAST / Echo in condition
      var fastMatch = cond.match(/(?:FAST|echo|POCUS)[^.]*(?::|—)\s*([^.]+\.)/i);
      if (fastMatch) results.push({label: 'Echo / POCUS · Stage ' + stageIdx, text: fastMatch[1].trim()});

      // IAP
      var iapMatch = (vitals + ' ' + cond).match(/IAP[^,\n]*/i);
      if (iapMatch) results.push({label: 'IAP · Stage ' + stageIdx, text: iapMatch[0].trim()});

      return results;
    }

    function buildCards() {
      var container = document.getElementById('inv-cards');
      var rn = document.getElementById('inv-room-note');
      if (rn) rn.textContent = 'This case projects to: ' + (DISPLAY_ROOM==='main' ? 'Main Hall display' : 'Room ' + DISPLAY_ROOM.toUpperCase() + ' display') + ' (/display?room=' + DISPLAY_ROOM + ')';
      var all = [];
      STAGES.forEach(function(s, i) {
        all = all.concat(parseInvestigations(s, i+1));
      });
      if (!all.length) { container.innerHTML = '<div style="font-size:0.75rem;color:#334155;padding:4px 0;">No pre-scripted results detected.</div>'; return; }
      container.innerHTML = all.map(function(inv, idx) {
        return '<button class="inv-card-btn" id="icard-' + idx + '" onclick="pushHtml(' + JSON.stringify(inv.label) + ',' + JSON.stringify(inv.text) + ',this)">'
          + '<span class="inv-card-label">' + inv.label + '</span>'
          + '<span class="inv-card-text">' + inv.text.substring(0,60) + (inv.text.length>60?'…':'') + '</span>'
          + '</button>';
      }).join('');
    }

    function pushCustom() {
      var result = document.getElementById('inv-text').value.trim();
      if (!result) { document.getElementById('inv-status').textContent = 'Enter a result first.'; return; }
      pushHtml('Investigation Result', result, document.getElementById('inv-push'));
    }

    function pushMedia(url, label, kind, btnEl) {
      var t = ts();
      var media = kind === 'video'
        ? '<video src="' + url + '" autoplay loop muted playsinline controls style="max-width:100%;max-height:78vh;border-radius:6px;box-shadow:0 4px 30px rgba(0,0,0,0.5);"></video>'
        : '<img src="' + url + '" style="max-width:100%;max-height:78vh;border-radius:6px;box-shadow:0 4px 30px rgba(0,0,0,0.5);" />';
      var html = '<div style="padding:24px 40px;display:flex;flex-direction:column;align-items:center;">'
        + '<div style="width:100%;display:flex;align-items:baseline;gap:16px;margin-bottom:20px;border-bottom:1px solid rgba(255,255,255,0.06);padding-bottom:14px;">'
        + '<div style="font-size:2rem;font-weight:800;color:#f1f5f9;letter-spacing:-0.02em;">' + label + '</div>'
        + '<div style="font-size:0.9rem;color:#334155;font-family:\'Courier New\',monospace;margin-left:auto;">' + t + '</div>'
        + '</div>'
        + media
        + '<div style="margin-top:20px;font-size:0.68rem;color:#1e3a5f;letter-spacing:0.14em;text-transform:uppercase;align-self:flex-start;">SimWars 2026 · Case ' + CASE_ID + '</div>'
        + '</div>';
      fetch('/api/score', {method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({key:'proj_html_' + DISPLAY_ROOM, value: html})})
        .then(function() {
          if (btnEl) { btnEl.style.borderColor = '#16a34a'; setTimeout(function(){ btnEl.style.borderColor = ''; }, 3500); }
          document.getElementById('inv-status').textContent = '✓ ' + label + ' pushed at ' + t + ' → Display ' + DISPLAY_ROOM.toUpperCase();
        })
        .catch(function() { document.getElementById('inv-status').textContent = 'Error — check connection.'; });
    }

    buildCards();
  </script>
</body>
</html>
"""

# ============================================
# ROUTES
# ============================================

@app.route('/')
def root_redirect():
    return redirect('/register')

@app.route('/library')
def index():
    locked = require_role()
    if locked: return locked
    rounds = {}
    active_cases = [c for c in CASES if c.get('active', True)]
    for c in active_cases:
        rounds.setdefault(c['round'], []).append(c)
    return render_template_string(INDEX_TEMPLATE, cases=active_cases, rounds=rounds)

IMG_EXTS = ('.jpg', '.jpeg', '.png', '.webp', '.gif')
VID_EXTS = ('.mp4', '.mov', '.webm', '.m4v')

def get_case_images(case_id):
    """Auto-discover real investigation images (CXR / Echo-POCUS / CT) dropped into
    static/investigations/<CASE_ID>/. Name files like '01_cxr-admission.jpg' —
    the leading number controls display order, the rest becomes the label.
    Nothing to configure in code: drop a file in the folder and it appears here."""
    folder = os.path.join(os.path.dirname(__file__), 'static', 'investigations', case_id)
    if not os.path.isdir(folder):
        return []
    images = []
    for f in sorted(os.listdir(folder)):
        if not f.lower().endswith(IMG_EXTS + VID_EXTS):
            continue
        stem = os.path.splitext(f)[0]
        label = re.sub(r'^\d+[_-]\s*', '', stem).replace('_', ' ').replace('-', ' ').strip()
        label = label.title() if label else f
        images.append({'file': f, 'label': label, 'url': f'/static/investigations/{case_id}/{f}',
                       'kind': 'video' if f.lower().endswith(VID_EXTS) else 'image'})
    return images

@app.route('/case/<case_id>')
def get_case(case_id):
    locked = require_role()
    if locked: return locked
    case = next((c for c in CASES if c['id'] == case_id), None)
    if not case:
        return "Case not found", 404

    fmt = request.args.get('format', 'html').lower()
    colors = ROUND_COLORS.get(case['round'], ROUND_COLORS['BLS CASES'])

    # Prev/next navigation
    ids = [c['id'] for c in CASES]
    idx = ids.index(case_id)
    prev_case = CASES[idx - 1] if idx > 0 else None
    next_case = CASES[idx + 1] if idx < len(CASES) - 1 else None

    if fmt == 'html':
        images = get_case_images(case_id)
        html = render_template_string(CASE_TEMPLATE, case=case, colors=colors, prev_case=prev_case, next_case=next_case, images=images, role=session.get('role'))
        return html, 200, {'Content-Type': 'text/html'}

    # PDF/DOCX exports contain the full director script and answer key — organiser only.
    if session.get('role') != 'organiser':
        return "Full case export is organiser-only. Judges can view the case summary at /case/" + case_id, 403

    elif fmt == 'pdf':
        WeasyHTML = _get_weasyprint()
        html = render_template_string(CASE_TEMPLATE, case=case, colors=colors, prev_case=None, next_case=None)
        if WeasyHTML is None:
            # No PDF engine on the server: serve a print-ready page that opens the
            # browser's own Save-as-PDF dialog instead of failing with a 503.
            html = html.replace('</body>', '<script>window.addEventListener("load",function(){setTimeout(function(){window.print();},400);});</scr' + 'ipt></body>')
            return html, 200, {'Content-Type': 'text/html'}
        pdf = WeasyHTML(string=html).write_pdf()
        return send_file(io.BytesIO(pdf), mimetype='application/pdf', as_attachment=True,
                         download_name=f'{case_id}_{case["title"][:40].replace(" ", "_")}.pdf')

    elif fmt == 'docx':
        doc = Document()
        doc.add_heading(case['title'], 0)
        doc.add_paragraph(case['round'], style='Intense Quote')
        doc.add_heading('Case Summary', level=1)
        doc.add_paragraph(case['summary'])
        doc.add_heading('Background (for participants)', level=1)
        doc.add_paragraph(case['background'])
        doc.add_heading('Expanded Background (if asked)', level=1)
        doc.add_paragraph(case['expanded_history'])
        doc.add_heading('Scenario Stages', level=1)
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = 'Table Grid'
        hdr = tbl.rows[0].cells
        for i, h in enumerate(['Stage', 'Vitals', 'Condition', 'Expected Interventions', 'Operator Notes']):
            hdr[i].text = h
        for s in case['stages']:
            row = tbl.add_row().cells
            row[0].text = s['name']
            row[1].text = s['vitals']
            row[2].text = s['condition']
            row[3].text = s['expected']
            row[4].text = s['notes']
        doc.add_heading('Patient Chart', level=1)
        p = doc.add_paragraph()
        for k, v in [('Name', case['patient']['name']), ('MRN', case['patient']['mrn']),
                     ('Gender', case['patient']['gender']), ('Age', case['patient']['age']),
                     ('DOB', case['patient']['dob']), ('Height', case['patient']['height']),
                     ('Weight', case['patient']['weight']), ('Allergies', case['patient']['allergies']),
                     ('Chief Complaint', case['patient']['cc']), ('HPI', case['patient']['hpi']),
                     ('PMH', case['patient']['pmh']), ('PSH', case['patient']['psh']),
                     ('Medications', case['patient']['meds']), ('Family/Social', case['patient']['family'])]:
            run = p.add_run(f'{k}: ')
            run.bold = True
            p.add_run(f'{v}\n')
        doc.add_heading('Actor Roles', level=1)
        doc.add_paragraph(case['actors'])
        doc.add_heading('Equipment', level=1)
        doc.add_paragraph(case['equipment'])
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True, download_name=f'{case_id}_{case["title"][:40].replace(" ", "_")}.docx')

    return "Invalid format. Use ?format=html, pdf, or docx", 400

@app.route('/scoring')
def scoring():
    locked = require_role()
    if locked: return locked
    # Inject the server-verified role + claimed judge slot so the page cannot be
    # tricked into organiser view via a URL parameter.
    path = os.path.join(os.path.dirname(__file__), 'scoring.html')
    with open(path, 'r', encoding='utf-8') as f:
        html = f.read()
    inject = ('<script>window.SIMWARS_ROLE = %s; window.SIMWARS_JUDGE_SLOT = %s;</script>\n</head>'
              % (_json.dumps(session.get('role')), _json.dumps(session.get('judge_slot') or '')))
    html = html.replace('</head>', inject, 1)
    return html, 200, {'Content-Type': 'text/html'}

import json as _json

# Actual round/room -> case allocations. Kept server-side only; only ever sent to the
# browser when the visiting session is already authenticated as judge or organiser via
# /unlock. Participants get zero trace of this in the HTML — not just CSS-hidden.
# NOTE: BLS (Room B) order below follows the LIVE scoring.html schedule
# (B1->Round1, B4->Round2, B3->Round3, B2->Round4). A separate reference file,
# simwars-2026-judges-master-schedule.html, states a different BLS order (B1,B2,B4,B3) —
# RESOLVED 15 Aug 2026: the LIVE order above is canonical. Every live surface (scoring
# engine, flow page, printable schedule) agrees on it; the reference file is outdated.
FLOW_CASE_DETAIL = {
    "prelim": {
        "1": {"A": {"code": "P1", "title": "Cardiogenic Shock — Fulminant Myocarditis",
                    "patient": "Sita · 1 yr · 10 kg · 3 days flu-like illness, cold shock",
                    "pearl": "No fluid boluses (the fluid trap) — Adrenaline inotrope; VT stays SYNCHRONISED 0.5 J/kg even when pulseless (organised rhythm)"},
              "B": {"code": "B2", "title": "Torsades Arrest — Congenital Long QT Syndrome",
                    "patient": "Varsha · 11 yr · 30 kg · Known LQTS, off beta-blockers",
                    "pearl": "Name Torsades, not just VT — Magnesium 50 mg/kg + unsynchronised 2 J/kg BOTH; avoid Amiodarone (prolongs QT)"}},
        "2": {"A": {"code": "P3", "title": "Acute Severe Asthma with Tension Pneumothorax",
                    "patient": "Arun · 3 yr · 14 kg · Steroid-hesitant mother",
                    "pearl": "Ketamine + Rocuronium RSI — post-intubation crash = right tension pneumothorax → needle decompression 2nd ICS MCL BEFORE any CXR"},
              "B": {"code": "B4", "title": "Refractory Status Epilepticus with Raised ICP",
                    "patient": "Sita · 12 yr · 30 kg · Unequal pupils, Cushing triad",
                    "pearl": "Unequal pupils in status = raised ICP — 3% NaCl / Mannitol + head up 30° + Ketamine RSI, never Propofol"}},
        "3": {"A": {"code": "P4", "title": "Refractory Septic Shock — Physiologically Difficult Airway",
                    "patient": "Maya · 3 yr · 15 kg · Meningococcal purpura, peri-arrest",
                    "pearl": "Resuscitate BEFORE intubating — vasopressor running + push-dose Epi 10 mcg/kg at the bedside before induction; Ketamine only"},
              "B": {"code": "B3", "title": "Electrocution — VF Arrest with Hyperkalaemia",
                    "patient": "Arush · 13 yr · 30 kg · K⁺ 7.5, entry/exit burns",
                    "pearl": "Defibrillate 2 J/kg immediately, then Calcium gluconate for K⁺ 7.5 — the critical missed step; avoid succinylcholine"}},
        "4": {"A": {"code": "P2", "title": "Septic Shock — Post-Open Heart Surgery",
                    "patient": "Manoj · 2 yr · 8 kg · 3 weeks post-ASD repair",
                    "pearl": "Glucose 28 first; the wide-QRS rhythm is sinus tach from Na 122 — do NOT cardiovert; cultures then Meropenem + Vancomycin"},
              "B": {"code": "B1", "title": "Near Drowning — Submersion Arrest with Severe Hypothermia",
                    "patient": "Vaibhavi · 12 yr · 30 kg · T 33°C, fine VF, 15 min CPR",
                    "pearl": "Not dead until warm and dead — defibrillate FIRST (EMS already gave Adrenaline ×2); active rewarming; ECMO if refractory"}},
    },
    "sf": {
        "1": {"code": "SF1", "title": "The Quiet Head — Dual Paediatric Trauma",
              "patient": "Ananya Rao · 7 yr · 24 kg + Vihaana Rao · 18 mo · 10 kg",
              "pearl": "Lucid interval is a trap — the outside CT is genuinely normal (imaged too early); treat the GCS and pupil trend, not the scan; calibrated response to Vihaana’s trace FAST"},
        "2": {"code": "SF2", "title": "The Loud Wound — Dual Paediatric Trauma",
              "patient": "Meera Iyer · 8 yr · 26 kg + Diya Iyer · 3 yr · 14 kg",
              "pearl": "Loud visible bleeding (Diya) is the distractor — positive FAST + shock (Meera) = theatre/IR not CT; activate MTP 1:1:1 + TXA"},
    },
    "final": {
        "code": "F2", "title": "The Tense Abdomen — Occult ACS behind a Contested-Authority, Septic-Shock Anchor",
        "patient": "Rehan · 10 yr · 30 kg · Day 5 undifferentiated febrile illness · ~70mL/kg crystalloid pre-loaded",
        "pearl": "Measure IAP (26 — Grade IV ACS): STOP fluids, therapeutic paracentesis, hold through the post-drain dip. Challenge Dr Sindhu’s half-right septic plan with evidence (IAP trend, urine output, dengue IgM) — diagnosis is never named; NS1 is genuinely negative",
    },
}

@app.route('/flow')
def flow():
    # Public route — participants can view without logging in and see only the sealed /
    # color-coded schedule. Judges/organisers who are already logged in (via /unlock) get
    # the actual case allocations injected server-side. Unauthenticated visitors receive
    # NO case-mapping data in the response at all (not merely hidden), so page source /
    # devtools reveals nothing extra — same standard as the projector display fix.
    path = os.path.join(os.path.dirname(__file__), 'flow-of-program.html')
    with open(path, 'r', encoding='utf-8') as f:
        html = f.read()
    role = session.get('role')
    if role in ('organiser', 'judge'):
        inject = (
            '<script>window.SIMWARS_ROLE = %s; window.SIMWARS_CASE_DETAIL = %s;</script>\n</head>'
            % (_json.dumps(role), _json.dumps(FLOW_CASE_DETAIL))
        )
    else:
        inject = '</head>'
    html = html.replace('</head>', inject, 1)
    return html, 200, {'Content-Type': 'text/html'}

@app.route('/schedule-print')
def schedule_print():
    return send_file(os.path.join(os.path.dirname(__file__), 'prelim-schedule-printable.html'))

@app.route('/register')
def register():
    # no-store: kills stale cached copies of the landing page on every device/QR link
    resp = send_file(os.path.join(os.path.dirname(__file__), 'simwars-2026-landing-page.html'))
    resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    resp.headers['Pragma'] = 'no-cache'
    return resp

@app.route('/version')
def version_stamp():
    # sanity check: open /version on the live site — must say BUILD-2026-08-16-D
    return 'BUILD-2026-08-16-D', 200, {'Content-Type': 'text/plain', 'Cache-Control': 'no-store'}

JUDGE_SLOTS = ('cj1', 'cj2', 'cj3', 'crm1', 'crm2', 'crm3')

def _key_visible_to_judge(key, my_slot, deb_revealed):
    """Judge-session score filtering. A judge only ever receives their own
    sheet entries; debriefer entries stay secret until revealed; prelim marks
    are split by room and domain. Organiser sessions are never filtered."""
    if key.startswith('sf_') or key.startswith('fn_'):
        # SF and Finals keys embed the judge slot (same six-judge panel).
        # Collective keys (penalties, shared, names) embed no slot and pass through.
        for s in JUDGE_SLOTS:
            if ('_%s_' % s) in key and s != my_slot:
                return False
    if key.startswith('deb_') and my_slot != 'debriefer' and not deb_revealed:
        return False
    # Prelim sheets: c_<CASE>_<TEAM>_* = Domain I (clinical), nt_... = Domain II (CRM)
    if key.startswith('c_') or key.startswith('nt_'):
        if my_slot == 'debriefer':
            return True
        parts = key.split('_')
        case = parts[1] if len(parts) > 1 else ''
        room = 'pals' if case.startswith('P') else ('bls' if case.startswith('B') else '')
        if my_slot in ('pals', 'bls'):
            return room == my_slot
        if my_slot.endswith('-cj'):
            return room == my_slot[:-3] and key.startswith('c_')
        if my_slot.endswith('-crm'):
            return room == my_slot[:-4] and key.startswith('nt_')
        return False
    return True

@app.route('/api/scores', methods=['GET'])
def get_scores():
    with get_db() as conn:
        rows = conn.execute('SELECT key, value FROM scores').fetchall()
    data = {row['key']: row['value'] for row in rows}
    role = session.get('role')
    if role not in ('organiser', 'judge', 'team'):
        # Public (projector displays): only the pushed display payloads, never scores.
        return jsonify({k: v for k, v in data.items() if k.startswith('proj_html')})
    if role == 'team':
        # Teams poll the scoreboard via /api/results; raw score keys stay sealed.
        return jsonify({k: v for k, v in data.items() if k.startswith('proj_html')})
    if role == 'judge':
        my_slot = session.get('judge_slot') or ''
        deb_revealed = data.get('debriefer_revealed') == 'true'
        data = {k: v for k, v in data.items() if _key_visible_to_judge(k, my_slot, deb_revealed)}
    return jsonify(data)

# One PIN per judge slot — the MC hands each judge their PIN on paper on event day.
# Change these before the event if they have been shared beyond the panel.
JUDGE_PINS = {
    'cj1': '7311', 'cj2': '7322', 'cj3': '7333', 'crm1': '7411', 'crm2': '7422',
    'crm3': '7433',
    'pals': '7100', 'pals-cj': '7101', 'pals-crm': '7102',
    'bls': '7200', 'bls-cj': '7201', 'bls-crm': '7202',
    'debriefer': '7500',
}

@app.route('/api/judge-identity', methods=['POST'])
def set_judge_identity():
    if session.get('role') not in ('judge', 'organiser'):
        return jsonify({'ok': False, 'error': 'not authenticated'}), 403
    data = request.get_json(silent=True) or {}
    slot = data.get('slot', '')
    if slot not in ('cj1', 'cj2', 'cj3', 'crm1', 'crm2', 'crm3', 'pals', 'pals-cj', 'pals-crm',
                    'bls', 'bls-cj', 'bls-crm', 'debriefer', ''):
        return jsonify({'ok': False, 'error': 'unknown slot'}), 400
    # Judges must present the slot's PIN; organiser sessions (the MC) never need one.
    if slot and session.get('role') == 'judge':
        if str(data.get('pin', '')) != JUDGE_PINS.get(slot):
            return jsonify({'ok': False, 'error': 'bad pin'}), 403
    session['judge_slot'] = slot
    return jsonify({'ok': True, 'slot': slot})

@app.route('/api/gate', methods=['POST'])
def gate():
    """Server-side password check for the landing page — keeps passwords out of
    the public page source and signs the session in at the same time."""
    pw = str((request.get_json(silent=True) or {}).get('password', '')).strip()
    if pw == ORGANISER_PASSWORD:
        session['role'] = 'organiser'
        return jsonify({'ok': True, 'role': 'organiser'})
    if pw == JUDGE_PASSWORD:
        session['role'] = 'judge'
        return jsonify({'ok': True, 'role': 'judge'})
    if pw.startswith(TEAM_PASSWORD_PREFIX):
        suffix = pw[len(TEAM_PASSWORD_PREFIX):]
        if suffix.isdigit() and 1 <= int(suffix) <= 16:
            session['role'] = 'team'
            session['team_number'] = int(suffix)
            return jsonify({'ok': True, 'role': 'team', 'team': int(suffix)})
    return jsonify({'ok': False}), 403

@app.route('/api/score', methods=['POST'])
def save_score():
    if session.get('role') not in ('organiser', 'judge'):
        return jsonify({'error': 'locked'}), 403
    data = request.get_json()
    key = data.get('key')
    value = data.get('value')
    if not key:
        return jsonify({'error': 'missing key'}), 400
    with get_db() as conn:
        conn.execute('INSERT INTO scores (key, value) VALUES (?, ?) ON CONFLICT(key) DO UPDATE SET value=excluded.value, updated_at=CURRENT_TIMESTAMP',
                     (key, str(value)))
        conn.commit()
    return jsonify({'ok': True})

@app.route('/api/scores/reset', methods=['POST'])
def reset_scores():
    if session.get('role') != 'organiser':
        return jsonify({'error': 'locked'}), 403
    with get_db() as conn:
        conn.execute('DELETE FROM scores')
        conn.commit()
    return jsonify({'ok': True})

@app.route('/api/results')
def get_results():
    with get_db() as conn:
        rows = conn.execute('SELECT key, value FROM scores').fetchall()
    sc = {row['key']: row['value'] for row in rows}
    def flt(val, default=0):
        try: return float(val) if val not in (None, '') else default
        except: return default
    ALL_TEAMS = ['T1','T2','T3','T4','T5','T6','T7','T8','T9','T10','T11','T12','T13','T14','T15','T16']
    teams = []
    for team in ALL_TEAMS:
        name = (sc.get(f'team_name_{team}') or '').strip() or team
        pals = min(100, flt(sc.get(f'score_pals_{team}')))
        bls  = min(100, flt(sc.get(f'score_bls_{team}')))
        pals_case = sc.get(f'team_pals_case_{team}') or ''
        bls_case  = sc.get(f'team_bls_case_{team}')  or ''
        pals_d3 = flt(sc.get(f'd3_{pals_case}_{team}')) if pals_case else 0
        bls_d3  = flt(sc.get(f'd3_{bls_case}_{team}'))  if bls_case  else 0
        combined = round(pals + bls, 1)
        teams.append({
            'team': team, 'name': name,
            'pals': pals, 'bls': bls, 'combined': combined,
            'pals_d3': pals_d3, 'bls_d3': bls_d3,
        })
    # Tie-breaker (locked): equal combined totals are split by the interprofessional/CRM
    # domain (Domain 3) summed across both prelim cases, then by the PALS score.
    teams.sort(key=lambda t: (-t['combined'], -(t['pals_d3'] + t['bls_d3']), -t['pals']))
    active = [t for t in teams if t['pals'] > 0 or t['bls'] > 0]
    sf = {t: (sc.get(f'sfName_{t}') or '').strip() or t.upper() for t in ['sf1','sf2','sf3','sf4']}
    fn_names = {t: (sc.get(f'fnName_{t}') or '').strip() or t.upper() for t in ['fin1','fin2']}
    FN_STAGES = {
        # Finals = single case F2 (revised). Judges score directly /200: Domain I Clinical 90,
        # Domain II Human Factors 80 (five 0-16 anchors), Domain III Shared Reasoning 30,
        # minus penalties. F1 is out of the competition.
        'F2': {'stages': [
            [('s1i0',2),('s1i1',2),('s1i2',4),('s1i3',2),('s1i4',4),('s1i5',10),('s1i6',2)],
            [('s2i0',4)],
            [('s3i0',10),('s3i1',8),('s3i2',8),('s3i3',2)],
            [('s4i0',12),('s4i1',6),('s4i2',6)],
            [('s5i0',2),('s5i1',6)],
        ], 'penalties': [('p0',-16),('p1',-10),('p2',-10),('p3',-8),('p4',-6),('p5',-6),('p6',-20)],
           'tw': [('tw0',16),('tw1',16),('tw2',16),('tw3',16),('tw4',16)]},
    }
    # Finals: six judges (same SF panel). Domain I = mean of three clinical sheets
    # + penalties (recorded once); Domain II = mean of three CRM sheets; Domain III collective.
    fn_scores = {}
    for t in ['fin1','fin2']:
        fn_scores[t] = {}
        for cid, cd in FN_STAGES.items():
            cj_totals = []
            for j in ('cj1', 'cj2', 'cj3'):
                tot = 0
                for stage in cd['stages']:
                    for k, mx in stage:
                        tot += min(mx, max(0, flt(sc.get(f'fn_{cid}_{t}_{j}_clin_{k}'))))
                cj_totals.append(tot)
            pen = 0
            for pk, pts in cd['penalties']:
                if sc.get(f'fn_{cid}_{t}_{pk}') in ('1', 1, True):
                    pen += pts
            clin = max(0, sum(cj_totals) / 3.0 + pen)
            crm_totals = [sum(min(mx, max(0, flt(sc.get(f'fn_{cid}_{t}_{j}_tw_{k}')))) for k, mx in cd['tw'])
                          for j in ('crm1', 'crm2', 'crm3')]
            tw = sum(crm_totals) / 3.0
            shared = min(30, max(0, flt(sc.get(f'fn_{cid}_{t}_shared'))))
            fn_scores[t][cid] = {'clin': round(clin, 1), 'tw': round(tw, 1), 'shared': shared,
                                 'total': round(clin + tw + shared, 1)}
    fn_totals = {t: fn_scores[t]['F2']['total'] for t in ['fin1','fin2']}  # judges score directly /200
    # Access tiers: organisers always; judges from 21 Aug IST; teams from 22 Aug IST.
    # Everyone else (incl. the public /flow page) gets team names only — no scores.
    role = session.get('role')
    allowed = role == 'organiser' \
        or (role == 'judge' and _ist_now() >= SCOREBOARD_OPENS['judge']) \
        or (role == 'team' and _ist_now() >= SCOREBOARD_OPENS['team'])
    if not allowed:
        return jsonify({'all_teams': [{'team': t['team'], 'name': t['name']} for t in teams],
                        'prelim': [], 'sf': {}, 'finals': {}, 'locked': True})
    return jsonify({
        'prelim': active,
        'all_teams': teams,
        'sf': sf,
        'finals': {t: {'name': fn_names[t], 'scores': fn_scores[t], 'total': fn_totals[t]} for t in ['fin1','fin2']},
        'weights': {'policy': 'BLS Domain I x1.25 / Domain II x0.8; PALS unweighted. Fixed — see scoring.html.', 'tiebreak': 'Ties on combined /200 broken by interprofessional (Domain 3) total across both cases, then PALS score.'},
    })

@app.route('/cases')
def list_cases():
    return jsonify([{'id': c['id'], 'title': c['title'], 'round': c['round']} for c in CASES])

SCOREBOARD_TEMPLATE = """<!DOCTYPE html>
<html lang="en"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1">
<title>SIM WARS 2026 · Score Board</title><style>
*{box-sizing:border-box;margin:0;padding:0;}
body{font-family:'Inter',-apple-system,Segoe UI,sans-serif;background:#f4f1f7;color:#2a0a44;padding:0 0 3rem;}
.bar{background:#2a0a44;color:#fff;padding:1rem 1.5rem;display:flex;align-items:center;justify-content:space-between;flex-wrap:wrap;gap:.6rem;}
.bar h1{font-size:1.15rem;font-weight:900;}
.bar .role{font-size:.72rem;font-weight:800;text-transform:uppercase;letter-spacing:.06em;background:rgba(255,255,255,.14);border:1px solid rgba(255,255,255,.25);padding:.35rem .9rem;border-radius:8px;}
.bar a{color:#f5a623;font-size:.82rem;font-weight:700;text-decoration:none;}
.wrap{max-width:900px;margin:1.6rem auto 0;padding:0 1rem;}
.card{background:#fff;border-radius:1.1rem;padding:1.4rem 1.5rem;box-shadow:0 4px 16px rgba(42,10,68,.06);margin-bottom:1.3rem;}
.card h2{font-size:1rem;font-weight:800;margin-bottom:.9rem;color:#2a0a44;}
table{width:100%;border-collapse:collapse;font-size:.9rem;}
th{text-align:left;font-size:.7rem;text-transform:uppercase;letter-spacing:.07em;color:#6a5a72;padding:.45rem .6rem;border-bottom:2px solid #eee0ea;}
td{padding:.55rem .6rem;border-bottom:1px solid #f0e8f0;}
tr.me td{background:#fdf1f6;font-weight:800;}
.rank{font-weight:900;color:#d81b7a;width:2.4rem;}
.muted{color:#6a5a72;font-size:.82rem;}
.empty{padding:1.2rem;text-align:center;color:#6a5a72;font-size:.9rem;}
</style></head><body>
<div class="bar"><h1>🏆 SIM WARS 2026 — Team &amp; Score Board</h1><div style="display:flex;gap:.8rem;align-items:center;"><span class="role">{{ role }}{% if team_number %} · Team {{ team_number }}{% endif %}</span><a href="/logout" style="color:#fff;opacity:.85;">Log out</a><a href="/register">← Site</a></div></div>
<div class="wrap">
<div class="card"><h2>Team Board — Colour Groups (from the draw)</h2><div id="groups" class="empty">Loading teams…</div></div>
<div class="card"><h2>Prelim Standings — combined /200</h2><div id="prelim" class="empty">Loading…</div></div>
<div class="card"><h2>Semi-Finalists</h2><div id="sf" class="empty">—</div></div>
<div class="card"><h2>Finals</h2><div id="fin" class="empty">—</div></div>
<p class="muted">Scores appear as they are locked by the judges and organisers. Refreshes automatically every 30 seconds.</p>
</div>
<script>
var MY_TEAM = {{ my_team_js }};
var IS_ORG = {{ is_org_js }};
var INTAKE = 'https://script.google.com/macros/s/AKfycbwyhhMFVFHo1IQCQw97Hzfv8PSWgXWgChtmbvqnXDtTA7pKExHsvBM3hL5rL279EjOM/exec';
var GROUPS = [
  {name:'RED',    draws:[1,2,3,4],     bg:'#fdecea', fg:'#b02a37'},
  {name:'GREEN',  draws:[5,6,7,8],     bg:'#e8f7f1', fg:'#0d8a72'},
  {name:'BLUE',   draws:[9,10,11,12],  bg:'#eef6fc', fg:'#2f6fd1'},
  {name:'YELLOW', draws:[13,14,15,16], bg:'#fef9c3', fg:'#a16207'}
];
function esc(s){return String(s==null?'':s).replace(/[&<>"']/g,function(c){return{'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'}[c];});}
function loadGroups(){
 fetch(INTAKE).then(function(r){return r.json();}).then(function(d){
  var byDraw={};(d.teams||[]).forEach(function(t){byDraw[parseInt(t.drawNumber,10)]=t;});
  var h='<div style="display:grid;grid-template-columns:repeat(auto-fit,minmax('+(IS_ORG?'260':'190')+'px,1fr));gap:.8rem;">';
  GROUPS.forEach(function(g){
   h+='<div style="background:'+g.bg+';border-radius:12px;padding:.8rem .9rem;">'
    +'<div style="font-size:.7rem;font-weight:900;letter-spacing:.1em;color:'+g.fg+';margin-bottom:.5rem;">'+g.name+' GROUP</div>'
    +g.draws.map(function(n){
      var t=byDraw[n];var me=MY_TEAM&&n===MY_TEAM;
      var head='<div style="font-size:.84rem;'+(me?'font-weight:800;':'')+'"><b style="color:'+g.fg+';">'+n+'</b> · '+(t?esc(t.teamName):'<span style="color:#9aa0ae;font-style:italic;">TBD</span>')+'</div>';
      var members='';
      if(IS_ORG&&t){
        members='<div style="margin-top:.3rem;font-size:.74rem;color:#5a5266;line-height:1.5;">'
          +[[t.m1name,t.m1desig],[t.m2name,t.m2desig],[t.m3name,t.m3desig],[t.m4name,t.m4desig]]
            .filter(function(m){return m[0];})
            .map(function(m){return esc(m[0])+(m[1]?' <span style="color:#9aa0ae;">· '+esc(m[1])+'</span>':'');})
            .join('<br>')
          +(t.leaderName?'<div style="margin-top:.2rem;color:'+g.fg+';font-weight:700;">Lead: '+esc(t.leaderName)+'</div>':'')
          +'</div>';
      }
      return '<div style="background:#fff;border-radius:8px;padding:.4rem .6rem;margin-bottom:.35rem;'+(me?'outline:2px solid '+g.fg+';':'')+'">'+head+members+'</div>';
     }).join('')+'</div>';
  });
  h+='</div>';
  var el=document.getElementById('groups');el.className='';el.innerHTML=h;
 }).catch(function(){document.getElementById('groups').textContent='Could not reach the registration sheet.';});
}
loadGroups();
function load(){
 fetch('/api/results').then(function(r){return r.json();}).then(function(d){
  var rows=(d.prelim&&d.prelim.length?d.prelim:[]);
  var el=document.getElementById('prelim');
  if(!rows.length){el.className='empty';el.textContent='No scores locked yet — check back after the first prelim rounds.';}
  else{
   var h='<table><tr><th></th><th>Team</th><th>PALS</th><th>BLS</th><th>Combined</th></tr>';
   rows.forEach(function(t,i){var me=MY_TEAM&&t.team==='T'+MY_TEAM;
    h+='<tr'+(me?' class="me"':'')+'><td class="rank">'+(i+1)+'</td><td>'+esc(t.team)+' — '+esc(t.name)+'</td><td>'+t.pals+'</td><td>'+t.bls+'</td><td><b>'+t.combined+'</b></td></tr>';});
   h+='</table>';el.className='';el.innerHTML=h;
  }
  var sf=d.sf||{};var sfv=Object.keys(sf).map(function(k){return sf[k];}).filter(function(v){return v&&!/^SF\\d$/i.test(v);});
  var sfe=document.getElementById('sf');
  if(sfv.length){sfe.className='';sfe.innerHTML=sfv.map(function(n){return '<span style="display:inline-block;background:#f0fbf9;color:#0d9488;font-weight:800;border-radius:8px;padding:.4rem .9rem;margin:.2rem;">'+esc(n)+'</span>';}).join('');}
  else{sfe.className='empty';sfe.textContent='Semi-finalists announced after prelims conclude.';}
  var fe=document.getElementById('fin');var f=d.finals||{};
  var named=['fin1','fin2'].filter(function(k){return f[k]&&f[k].name&&!/^FIN\\d$/i.test(f[k].name);});
  if(named.length){
   var h2='<table><tr><th>Finalist</th><th>Total</th></tr>';
   named.forEach(function(k){h2+='<tr><td>'+esc(f[k].name)+'</td><td><b>'+f[k].total+'</b></td></tr>';});
   fe.className='';fe.innerHTML=h2+'</table>';
  } else {fe.className='empty';fe.textContent='Finalists announced after the semi-finals.';}
 }).catch(function(){});
}
load();setInterval(load,30000);
</script></body></html>"""

SCOREBOARD_GATE_TEMPLATE = """<!DOCTYPE html>
<html><head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1"><title>SIM WARS 2026 · Score Board</title>
<style>body{font-family:-apple-system,Segoe UI,Roboto,sans-serif;background:#2a0a44;color:#fff;display:flex;align-items:center;justify-content:center;min-height:100vh;margin:0;}
.card{background:rgba(255,255,255,.07);border:1px solid rgba(255,255,255,.15);padding:2.4rem 2.2rem;border-radius:18px;max-width:400px;width:90%;text-align:center;}
h1{font-size:1.25rem;margin:0 0 .5rem;}p{font-size:.9rem;color:rgba(255,255,255,.75);margin:0 0 1.3rem;line-height:1.5;}
input{width:100%;padding:.8rem 1rem;border-radius:10px;border:none;font-size:1rem;margin-bottom:1rem;box-sizing:border-box;}
button{width:100%;background:#d81b7a;color:#fff;border:none;padding:.85rem;border-radius:10px;font-weight:800;font-size:1rem;cursor:pointer;}
.err{background:rgba(216,27,122,.25);border:1px solid #d81b7a;border-radius:8px;padding:.6rem .8rem;font-size:.85rem;margin-bottom:1rem;}
.note{margin-top:1.1rem;font-size:.78rem;color:rgba(255,255,255,.55);line-height:1.5;}</style></head><body>
<form class="card" method="post">
<h1>🏆 Team &amp; Score Board</h1>
<p>{{ message }}</p>
{% if error %}<div class="err">{{ error }}</div>{% endif %}
{% if show_form %}<input type="password" name="password" placeholder="Your access password" autofocus autocomplete="off"><button type="submit">View Score Board</button>{% endif %}
<div class="note">Organisers: live now · Judges: from 21 Aug 2026 · Teams: on event day (22 Aug), password ends with your draw number</div>
</form></body></html>"""

@app.route('/scoreboard', methods=['GET', 'POST'])
def scoreboard():
    if request.method == 'POST':
        pw = request.form.get('password', '').strip()
        if pw == ORGANISER_PASSWORD:
            session['role'] = 'organiser'
        elif pw == JUDGE_PASSWORD:
            session['role'] = 'judge'
        elif pw.startswith(TEAM_PASSWORD_PREFIX) and pw[len(TEAM_PASSWORD_PREFIX):].isdigit() \
                and 1 <= int(pw[len(TEAM_PASSWORD_PREFIX):]) <= 16:
            session['role'] = 'team'
            session['team_number'] = int(pw[len(TEAM_PASSWORD_PREFIX):])
        else:
            return render_template_string(SCOREBOARD_GATE_TEMPLATE, show_form=True,
                message='Enter your access password to view live standings.',
                error='Incorrect password. Try again.')
    role = session.get('role')
    if role not in ('organiser', 'judge', 'team'):
        return render_template_string(SCOREBOARD_GATE_TEMPLATE, show_form=True,
            message='Enter your access password to view live standings.', error=None)
    opens = SCOREBOARD_OPENS.get(role)
    if opens and _ist_now() < opens:
        return render_template_string(SCOREBOARD_GATE_TEMPLATE, show_form=False,
            message=('The score board opens for %ss on %s IST. Come back then!'
                     % (role, opens.strftime('%d %b %Y, %H:%M'))), error=None), 403
    tn = session.get('team_number') if role == 'team' else None
    return render_template_string(SCOREBOARD_TEMPLATE, role=role, team_number=tn,
                                  my_team_js=(str(tn) if tn else 'null'),
                                  is_org_js=('true' if role == 'organiser' else 'false'))


@app.route('/logout')
def logout():
    # Full sign-out for every role (organiser, judge, team): clear the session
    # cookie and land on the neutral passcode gate.
    session.clear()
    return redirect('/scoreboard')


DISPLAY_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>SimWars · Investigations Display</title>
<style>
* { box-sizing: border-box; margin: 0; padding: 0; }
html, body { height: 100%; background: #030712; color: #e2e8f0; font-family: 'Courier New', Courier, monospace; }
#display-wrap { min-height: 100vh; padding: 40px 48px; }
#placeholder {
    display: flex; flex-direction: column; align-items: center; justify-content: center;
    min-height: 80vh; gap: 20px; color: #1e3a5f;
}
#placeholder-icon { font-size: 5rem; }
#placeholder-text { font-size: 1.3rem; letter-spacing: 0.05em; }
#placeholder-sub { font-size: 0.85rem; color: #0f172a; }
</style>
</head>
<body>
<div id="display-wrap">
    <div id="content">
        <div id="placeholder">
            <div id="placeholder-icon">📋</div>
            <div id="placeholder-text">Awaiting investigation result…</div>
            <div id="placeholder-sub">SimWars 2026 — Operator will project results here</div>
        </div>
    </div>
</div>
<script>
var ROOM = (new URLSearchParams(location.search).get('room') || '').toLowerCase();
if (ROOM !== 'a' && ROOM !== 'b' && ROOM !== 'main') {
    document.getElementById('content').innerHTML =
        '<div style="display:flex;flex-direction:column;align-items:center;justify-content:center;min-height:80vh;gap:26px;">'
        + '<div style="font-size:1.4rem;color:#64748b;letter-spacing:.08em;">SELECT THIS ROOM</div>'
        + ['a','b','main'].map(function(r){
            var lbl = r==='main' ? 'MAIN HALL · Semis & Finals' : 'ROOM ' + r.toUpperCase() + (r==='a' ? ' · PALS' : ' · BLS');
            return '<a href="/display?room=' + r + '" style="display:block;background:#0f172a;border:2px solid #1e3a5f;border-radius:16px;padding:26px 60px;color:#e2e8f0;text-decoration:none;font-size:1.6rem;letter-spacing:.05em;">' + lbl + '</a>';
          }).join('')
        + '</div>';
} else {
    document.getElementById('placeholder-sub').textContent = 'SimWars 2026 — ' + (ROOM==='main' ? 'Main Hall' : 'Room ' + ROOM.toUpperCase()) + ' · Operator will project results here';
    var _last = null;
    var KEY = 'proj_html_' + ROOM;
    function poll() {
        fetch('/api/scores').then(function(r){ return r.json(); }).then(function(data){
            var html = data[KEY] || data['proj_html'];
            if (html && html !== _last) {
                _last = html;
                document.getElementById('content').innerHTML =
                    '<div style="font-family:\\'Courier New\\',monospace;color:#e2e8f0;padding:8px 0;">' + html + '</div>';
            }
        }).catch(function(){}).finally(function(){ setTimeout(poll, 1500); });
    }
    poll();
}
</script>
</body>
</html>"""

@app.route('/display')
def display():
    return DISPLAY_TEMPLATE

PORTAL_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8" />
    <meta name="viewport" content="width=device-width, initial-scale=1.0" />
    <title>SIMWARS 2026 · PediSTARS India</title>
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0-beta3/css/all.min.css" />
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }

        body {
            font-family: 'Inter', 'Segoe UI', system-ui, sans-serif;
            background: #eef2f7;
            min-height: 100vh;
            padding: 2rem 1.5rem;
            background-image: radial-gradient(circle at 20% 30%, rgba(249,199,79,0.04) 0%, transparent 60%);
            display: flex;
            justify-content: center;
            align-items: flex-start;
        }

        .portal {
            max-width: 1300px;
            width: 100%;
            background: rgba(255,255,255,0.7);
            backdrop-filter: blur(6px);
            border-radius: 2.8rem;
            box-shadow: 0 40px 80px rgba(11,42,79,0.12), 0 10px 30px rgba(0,0,0,0.03);
            border: 1px solid rgba(255,255,255,0.6);
            overflow: hidden;
        }

        /* ── HERO ── */
        .hero {
            background: linear-gradient(135deg, #0a1a2e 0%, #1b3f64 70%, #2b5780 100%);
            padding: 2.2rem 3rem 1.8rem;
            color: white;
            position: relative;
            overflow: hidden;
        }

        .hero::before {
            content: '';
            position: absolute; top: -60%; right: -10%;
            width: 500px; height: 500px;
            background: radial-gradient(circle, rgba(249,199,79,0.07) 0%, transparent 70%);
            border-radius: 50%; pointer-events: none;
        }

        .hero::after {
            content: '';
            position: absolute; bottom: -50%; left: -5%;
            width: 350px; height: 350px;
            background: radial-gradient(circle, rgba(255,255,255,0.02) 0%, transparent 70%);
            border-radius: 50%; pointer-events: none;
        }

        .hero-top {
            display: flex; flex-wrap: wrap;
            justify-content: space-between; align-items: center;
            gap: 1.5rem; position: relative; z-index: 2;
        }

        .brand { display: flex; align-items: center; gap: 1.2rem; }

        .brand-icon {
            background: #f9c74f; color: #0a1a2e;
            width: 70px; height: 70px; border-radius: 20px;
            display: flex; align-items: center; justify-content: center;
            font-weight: 800; font-size: 1.8rem; letter-spacing: -0.02em;
            box-shadow: 0 10px 24px rgba(0,0,0,0.25); flex-shrink: 0;
        }

        .brand-text { line-height: 1.2; }
        .brand-text .name { font-size: 2rem; font-weight: 700; letter-spacing: -0.01em; }
        .brand-text .name span { color: #f9c74f; }
        .brand-text .tagline { font-size: 0.85rem; font-weight: 400; opacity: 0.7; letter-spacing: 0.25em; margin-top: 0.1rem; }

        .badge-location {
            background: rgba(255,255,255,0.1); backdrop-filter: blur(4px);
            padding: 0.4rem 1.6rem; border-radius: 60px;
            font-weight: 500; font-size: 0.85rem;
            border: 1px solid rgba(255,255,255,0.15); letter-spacing: 0.05em; white-space: nowrap;
        }

        .badge-location i { margin-right: 8px; color: #f9c74f; }

        .hero-title { margin-top: 2rem; position: relative; z-index: 2; }

        .hero-title h1 {
            font-size: 2.6rem; font-weight: 700;
            line-height: 1.15; letter-spacing: -0.02em;
        }

        .hero-title h1 .gold { color: #f9c74f; }

        .hero-title .sub {
            font-size: 1rem; opacity: 0.8;
            margin-top: 0.3rem; font-weight: 300; letter-spacing: 0.03em;
        }

        .hero-meta {
            display: flex; flex-wrap: wrap; gap: 1.8rem 3rem;
            margin-top: 1.6rem; padding-top: 1.2rem;
            border-top: 1px solid rgba(255,255,255,0.08);
            position: relative; z-index: 2;
        }

        .meta-item { display: flex; align-items: center; gap: 0.7rem; font-size: 0.9rem; }
        .meta-item i { color: #f9c74f; font-size: 1rem; width: 1.4rem; text-align: center; }
        .meta-item strong { font-weight: 600; }
        .meta-item .highlight { color: #f9c74f; }

        /* ── BODY ── */
        .body { padding: 2.5rem 3rem 1.8rem; }

        .grid {
            display: grid;
            grid-template-columns: 1.2fr 1fr;
            gap: 2.5rem;
        }

        @media(max-width: 900px) {
            .grid { grid-template-columns: 1fr; gap: 2rem; }
            .hero { padding: 1.8rem; }
            .hero-title h1 { font-size: 1.8rem; }
            .brand-text .name { font-size: 1.5rem; }
            .brand-icon { width: 54px; height: 54px; font-size: 1.4rem; }
            .body { padding: 1.8rem; }
        }

        /* ── SECTION HEADER ── */
        .section-head {
            font-size: 1.1rem; font-weight: 600; color: #0a1a2e;
            margin-bottom: 1.2rem;
            display: flex; align-items: center; gap: 0.6rem;
        }
        .section-head i { color: #f9c74f; }

        .section-sub {
            font-size: 0.65rem; font-weight: 700; letter-spacing: 0.14em;
            text-transform: uppercase; color: #5a7a9a;
            margin: 1.2rem 0 0.6rem;
        }

        /* ── FEATURE GRID ── */
        .feature-grid {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 1rem;
        }

        @media(max-width: 600px) { .feature-grid { grid-template-columns: 1fr; } }

        .feature-card {
            background: white; border-radius: 1.2rem;
            padding: 1.2rem 1.2rem 1rem;
            box-shadow: 0 4px 12px rgba(0,0,0,0.02);
            border: 1px solid rgba(0,0,0,0.02);
            transition: 0.25s ease;
            text-decoration: none; color: #0a1a2e;
            display: flex; flex-direction: column;
        }

        .feature-card:hover {
            transform: translateY(-4px);
            box-shadow: 0 12px 32px rgba(11,42,79,0.08);
            border-color: #dce5f0;
        }

        .feature-card .icon { font-size: 1.5rem; color: #f9c74f; margin-bottom: 0.5rem; }
        .feature-card .title { font-weight: 600; font-size: 0.95rem; margin-bottom: 0.2rem; }
        .feature-card .desc { font-size: 0.78rem; color: #5a7a9a; line-height: 1.4; flex: 1; }
        .feature-card .url {
            font-size: 0.62rem; color: #8aaac0; margin-top: 0.6rem;
            font-family: 'SF Mono', monospace;
            border-top: 1px solid #eef4fa; padding-top: 0.5rem;
        }

        .badge {
            display: inline-block;
            background: #f9c74f; color: #0a1a2e;
            font-size: 0.55rem; font-weight: 700;
            padding: 0.15rem 0.6rem; border-radius: 40px;
            text-transform: uppercase; letter-spacing: 0.04em;
            margin-left: 0.4rem;
        }

        .badge-live  { background: #22c55e; color: #fff; }
        .badge-bls   { background: #16a34a; color: #fff; }
        .badge-pals  { background: #1d4ed8; color: #fff; }
        .badge-sf    { background: #0891b2; color: #fff; }
        .badge-fn    { background: #d97706; color: #fff; }
        .badge-spec  { background: #7c3aed; color: #fff; }

        /* ── INFO PANEL (right column) ── */
        .info-panel { display: flex; flex-direction: column; gap: 1.2rem; }

        .info-card {
            background: white; border-radius: 1.5rem;
            padding: 1.4rem 1.8rem;
            box-shadow: 0 4px 12px rgba(0,0,0,0.02);
            border-left: 4px solid #f9c74f;
        }

        .info-card .label {
            font-size: 0.7rem; text-transform: uppercase;
            letter-spacing: 0.08em; color: #5a7a9a; font-weight: 600;
        }

        .info-card .value {
            font-size: 0.92rem; color: #0a1a2e;
            margin-top: 0.3rem; line-height: 1.7;
        }

        .info-card .value .coord {
            display: block; font-size: 0.85rem; color: #1a3a5a; margin-top: 0.1rem;
        }
        .info-card .value .coord i { color: #f9c74f; width: 1.2rem; }

        /* quick links inside info-card */
        .quick-links { display: flex; flex-direction: column; gap: 0.5rem; margin-top: 0.4rem; }

        .quick-link {
            display: flex; align-items: center; gap: 0.8rem;
            background: #f8faff; padding: 0.55rem 1rem;
            border-radius: 0.9rem; text-decoration: none;
            color: #0a1a2e; transition: 0.2s;
            border: 1px solid transparent; font-size: 0.88rem;
        }

        .quick-link:hover { background: white; border-color: #dce5f0; box-shadow: 0 4px 12px rgba(0,0,0,0.03); }
        .quick-link i { color: #f9c74f; width: 1.2rem; font-size: 0.95rem; }
        .quick-link .link-label { font-weight: 500; flex: 1; }
        .quick-link .link-badge { font-size: 0.58rem; }

        /* ── CASES SECTION (full width below grid) ── */
        .cases-section {
            margin-top: 2.5rem;
            padding-top: 2rem;
            border-top: 1px solid #e4ecf5;
        }

        .round-group { margin-bottom: 1.8rem; }

        .round-label {
            display: flex; align-items: center; gap: 8px;
            font-size: 0.65rem; font-weight: 700;
            letter-spacing: 0.16em; text-transform: uppercase;
            color: #5a7a9a; margin-bottom: 0.8rem;
        }

        .round-label .dot {
            width: 8px; height: 8px; border-radius: 50%; flex-shrink: 0;
        }

        .round-label::after { content: ''; flex: 1; height: 1px; background: #e4ecf5; }

        .dot-bls  { background: #22c55e; }
        .dot-pals { background: #3b82f6; }
        .dot-sf   { background: #06b6d4; }
        .dot-fn   { background: #f59e0b; }

        .case-grid {
            display: grid;
            grid-template-columns: repeat(4, 1fr);
            gap: 10px;
        }

        .case-grid-2 { grid-template-columns: 1fr 1fr; }

        .case-card {
            background: white; border-radius: 1rem;
            padding: 1rem 1.1rem;
            box-shadow: 0 2px 8px rgba(0,0,0,0.025);
            border: 1px solid #edf2fa;
            transition: 0.18s;
        }

        .case-card:hover { box-shadow: 0 6px 20px rgba(11,42,79,0.07); transform: translateY(-1px); }

        .round-bls  .case-card { border-top: 3px solid #22c55e; }
        .round-pals .case-card { border-top: 3px solid #3b82f6; }
        .round-sf   .case-card { border-top: 3px solid #06b6d4; }
        .round-fn   .case-card { border-top: 3px solid #f59e0b; }

        .case-id {
            font-size: 0.95rem; font-weight: 800;
            color: #0a1a2e; margin-bottom: 0.6rem; line-height: 1.3;
        }

        .doc-links { display: flex; flex-direction: column; gap: 4px; }

        .doc-link {
            display: inline-flex; align-items: center; gap: 5px;
            font-size: 0.68rem; font-weight: 500;
            color: #5a7a9a; text-decoration: none; padding: 2px 0;
            transition: color 0.15s;
        }

        .doc-link:hover { color: #0a1a2e; }

        /* ── FOOTER ── */
        .footer {
            text-align: center; font-size: 0.8rem; color: #6a8aaa;
            border-top: 1px solid #eef4fa;
            padding: 1.5rem 2rem 1.8rem; letter-spacing: 0.02em;
            background: rgba(255,255,255,0.3);
        }

        .footer i { color: #f9c74f; margin: 0 4px; }
        .footer .sep { opacity: 0.3; margin: 0 0.6rem; }

        @media(max-width: 580px) {
            .case-grid { grid-template-columns: 1fr 1fr; }
            .case-grid-2 { grid-template-columns: 1fr; }
            .hero-top { flex-direction: column; align-items: flex-start; }
        }
    </style>
</head>
<body>

<div class="portal">

    <!-- ══ HERO ══ -->
    <header class="hero">
        <div class="hero-top">
            <div class="brand">
                <div class="brand-icon">PS</div>
                <div class="brand-text">
                    <div class="name">Pedi<span>STARS</span></div>
                    <div class="tagline">INDIA · SIMULATION &amp; RESEARCH</div>
                </div>
            </div>
            <div class="badge-location">
                <i class="fas fa-map-pin"></i> Bengaluru · 22–23 Aug 2026
            </div>
        </div>

        <div class="hero-title">
            <h1>Simulus <span class="gold">11th Annual</span><br>National Conference of PediSTARS</h1>
            <div class="sub"><i class="fas fa-stethoscope" style="margin-right:10px"></i>Healthcare Simulation · SIMWARS 2026</div>
        </div>

        <div class="hero-meta">
            <div class="meta-item">
                <i class="fas fa-calendar-alt"></i>
                <span><strong>23 August 2026</strong> · Main Conference</span>
            </div>
            <div class="meta-item">
                <i class="fas fa-location-dot"></i>
                <span><strong>22 Aug</strong> Prelims @ DHEE Hospitals · <strong>23 Aug</strong> Semis &amp; Finals @ VASA</span>
            </div>
            <div class="meta-item">
                <i class="fas fa-users"></i>
                <span>Max <strong>16 teams</strong> · <span class="highlight">4 members</span> each</span>
            </div>
            <div class="meta-item">
                <i class="fas fa-tag"></i>
                <span>Registration <strong>₹10,000</strong> / <strong>$215</strong> per team</span>
            </div>
        </div>
    </header>

    <!-- ══ BODY ══ -->
    <div class="body">
        <div class="grid">

            <!-- LEFT: Conference tools + scoring sheets -->
            <div>
                <div class="section-head"><i class="fas fa-bolt"></i> Live competition tools</div>
                <div class="feature-grid">

                    <a href="https://merry-intuition-production-1788.up.railway.app/" target="_blank" class="feature-card">
                        <div class="icon"><i class="fas fa-book-open"></i></div>
                        <div class="title">Case Library <span class="badge badge-live">Live</span></div>
                        <div class="desc">All 14 cases · stage-by-stage breakdowns · investigations console · prelim standings.</div>
                        <div class="url">merry-intuition-production…/</div>
                    </a>

                    <a href="https://merry-intuition-production-1788.up.railway.app/scoring" target="_blank" class="feature-card">
                        <div class="icon"><i class="fas fa-bullseye"></i></div>
                        <div class="title">Scoring Engine <span class="badge badge-live">Live</span></div>
                        <div class="desc">Master · Judge · Debriefer · Special Prizes · Semis &amp; Finals. Live sync across devices.</div>
                        <div class="url">…/scoring</div>
                    </a>

                    <a href="https://merry-intuition-production-1788.up.railway.app/display" target="_blank" class="feature-card">
                        <div class="icon"><i class="fas fa-desktop"></i></div>
                        <div class="title">Results Display</div>
                        <div class="desc">Fullscreen investigations display for the room projector. Auto-updates live.</div>
                        <div class="url">…/display</div>
                    </a>

                    <a href="https://claude.ai/code/artifact/631528d3-55b7-45a7-b758-c2e6192fb7b9" target="_blank" class="feature-card">
                        <div class="icon"><i class="fas fa-diagram-project"></i></div>
                        <div class="title">Competition Flow</div>
                        <div class="desc">Full case descriptions · patient profiles · scoring domains · advancement rules.</div>
                        <div class="url">claude.ai/code/artifact/…</div>
                    </a>

                </div>

                <div class="section-sub">Scoring Sheets</div>
                <div class="feature-grid">

                    <a href="https://merry-intuition-production-1788.up.railway.app/scoring#tabBls" target="_blank" class="feature-card">
                        <div class="icon" style="color:#16a34a"><i class="fas fa-list-check"></i></div>
                        <div class="title">BLS Judge Sheet <span class="badge badge-bls">R1</span></div>
                        <div class="desc">BLS Round · Cases B1–B4 · Clinical + Non-technical rubric per team.</div>
                        <div class="url">…/scoring#tabBls</div>
                    </a>

                    <a href="https://merry-intuition-production-1788.up.railway.app/scoring#tabPals" target="_blank" class="feature-card">
                        <div class="icon" style="color:#1d4ed8"><i class="fas fa-list-check"></i></div>
                        <div class="title">PALS Judge Sheet <span class="badge badge-pals">R2</span></div>
                        <div class="desc">PALS Round · Cases P1–P4 · Clinical + Non-technical rubric per team.</div>
                        <div class="url">…/scoring#tabPals</div>
                    </a>

                    <a href="https://merry-intuition-production-1788.up.railway.app/scoring#tabSemiFinals" target="_blank" class="feature-card">
                        <div class="icon" style="color:#0891b2"><i class="fas fa-trophy"></i></div>
                        <div class="title">Semi-Finals Sheet <span class="badge badge-sf">SF</span></div>
                        <div class="desc">SF1 + SF2 · Domain I / II / III · Six judges + Chief · 40/50/10.</div>
                        <div class="url">…/scoring#tabSemiFinals</div>
                    </a>

                    <a href="https://merry-intuition-production-1788.up.railway.app/scoring#tabFinals" target="_blank" class="feature-card">
                        <div class="icon" style="color:#d97706"><i class="fas fa-medal"></i></div>
                        <div class="title">Finals Sheet <span class="badge badge-fn">Final</span></div>
                        <div class="desc">F2 single case · Six judges · Domains I / II / III · 90/80/30 · /200 total.</div>
                        <div class="url">…/scoring#tabFinals</div>
                    </a>

                    <a href="https://merry-intuition-production-1788.up.railway.app/scoring#tabMaster" target="_blank" class="feature-card">
                        <div class="icon"><i class="fas fa-grid-2"></i></div>
                        <div class="title">Master Sheet</div>
                        <div class="desc">Team setup · draw assignment · Grand Total across all 16 teams.</div>
                        <div class="url">…/scoring#tabMaster</div>
                    </a>

                    <a href="https://merry-intuition-production-1788.up.railway.app/scoring#tabSpecial" target="_blank" class="feature-card">
                        <div class="icon" style="color:#7c3aed"><i class="fas fa-star"></i></div>
                        <div class="title">Special Prizes <span class="badge badge-spec">6</span></div>
                        <div class="desc">Best CPR · CRM · Parental Comms · Interprofessional · Video · Uniform.</div>
                        <div class="url">…/scoring#tabSpecial</div>
                    </a>

                </div>
            </div>

            <!-- RIGHT: Info + quick links -->
            <div class="info-panel">

                <div class="info-card">
                    <div class="label">👥 National Coordinators</div>
                    <div class="value">
                        <strong>Dr Manju Kedarnath</strong>
                        <span class="coord"><i class="fas fa-phone-alt"></i> +91 99866 34300</span>
                        <strong style="display:block;margin-top:0.5rem">Dr Pritesh Nagar</strong>
                        <span class="coord"><i class="fas fa-phone-alt"></i> +91 99596 65000</span>
                        <strong style="display:block;margin-top:0.5rem">Dr Supraja Chandrasekhar</strong>
                        <span class="coord"><i class="fas fa-phone-alt"></i> +91 99458 18818</span>
                    </div>
                    <div style="margin-top:1rem">
                        <div class="label">📍 Local Coordinators</div>
                        <div class="value">
                            <strong>Dr Javed</strong>
                            <span class="coord"><i class="fas fa-phone-alt"></i> +91 81436 09989</span>
                            <strong style="display:block;margin-top:0.5rem">Ms Kavita Chandrakar</strong>
                            <span class="coord"><i class="fas fa-phone-alt"></i> +91 96858 63289</span>
                        </div>
                    </div>
                </div>

                <div class="info-card" style="border-left-color:#2b5780">
                    <div class="label">📌 Venues &amp; Dates</div>
                    <div class="value">
                        <strong>22 August</strong> — Preliminary Rounds<br>
                        DHEE Hospitals, Kanakapura Road, Bengaluru
                        <br><br>
                        <strong>23 August</strong> — Semi-Finals &amp; Finals<br>
                        VASA, Bengaluru
                        <br><br>
                        <span style="font-size:0.82rem;color:#2b5780">
                            <i class="fas fa-info-circle" style="color:#f9c74f"></i>
                            Max 16 teams · 4 members each · at least 2 doctors (PG trainees) · min 1 nurse
                        </span>
                    </div>
                </div>

                <div class="info-card" style="border-left-color:#a0c0d0">
                    <div class="label">🏆 Team Eligibility</div>
                    <div class="value" style="font-size:0.86rem">
                        <span class="coord" style="margin-top:0"><i class="fas fa-angle-right"></i> Physicians &amp; residents: &lt;5 yrs post-MD/DNB</span>
                        <span class="coord"><i class="fas fa-angle-right"></i> Nurses &amp; paramedics: no experience limit</span>
                        <span class="coord"><i class="fas fa-angle-right"></i> Members from different hospitals of same org allowed</span>
                        <span class="coord"><i class="fas fa-angle-right"></i> Max 1 person who competed in the previous SimWars</span>
                    </div>
                </div>

                <div class="info-card" style="border-left-color:#d4a017">
                    <div class="label">🗺️ Competition Reference</div>
                    <div class="quick-links">
                        <a href="https://claude.ai/code/artifact/12f7e01d-f7de-4d8a-b62a-1dd8af0439e1" target="_blank" class="quick-link">
                            <i class="fas fa-diagram-project"></i>
                            <span class="link-label">Flow — Full</span>
                            <span class="badge" style="margin-left:auto">All stages</span>
                        </a>
                        <a href="https://claude.ai/code/artifact/41b20ddd-fa61-4b86-8e57-4f81d7417ff9" target="_blank" class="quick-link">
                            <i class="fas fa-sitemap"></i>
                            <span class="link-label">Flow — Clean</span>
                            <span class="badge" style="margin-left:auto">Structure</span>
                        </a>
                        <a href="https://claude.ai/code/artifact/631528d3-55b7-45a7-b758-c2e6192fb7b9" target="_blank" class="quick-link">
                            <i class="fas fa-file-lines"></i>
                            <span class="link-label">Flow — Detailed</span>
                            <span class="badge" style="margin-left:auto">Full cases</span>
                        </a>
                    </div>
                </div>

            </div>
        </div>

        <!-- CASE PACKETS — full width -->
        <div class="cases-section">
            <div class="section-head"><i class="fas fa-folder-open"></i> Case Packets <span style="font-size:0.75rem;font-weight:400;color:#6a8aaa;margin-left:6px">Live server copies</span></div>

            <!-- BLS -->
            <div class="round-group round-bls">
                <div class="round-label"><div class="dot dot-bls"></div> BLS — Prelims Round 1</div>
                <div class="case-grid">
                    <div class="case-card">
                        <div class="case-id">B1</div>
                        <div class="doc-links">
                            <a class="doc-link" href="/case/B1" target="_blank">📄 Case Page</a>
                            <a class="doc-link" href="/case/B1?format=pdf" target="_blank">📋 Case PDF</a>
                            <a class="doc-link" href="/scoring" target="_blank">⚖️ Judge Scoring</a>
                            <a class="doc-link" href="/static/decks/B1_Slide_Deck.pptx" target="_blank">📊 Slide Deck</a>
                        </div>
                    </div>
                    <div class="case-card">
                        <div class="case-id">B2</div>
                        <div class="doc-links">
                            <a class="doc-link" href="/case/B2" target="_blank">📄 Case Page</a>
                            <a class="doc-link" href="/case/B2?format=pdf" target="_blank">📋 Case PDF</a>
                            <a class="doc-link" href="/scoring" target="_blank">⚖️ Judge Scoring</a>
                            <a class="doc-link" href="/static/decks/B2_Slide_Deck.pptx" target="_blank">📊 Slide Deck</a>
                        </div>
                    </div>
                    <div class="case-card">
                        <div class="case-id">B3</div>
                        <div class="doc-links">
                            <a class="doc-link" href="/case/B3" target="_blank">📄 Case Page</a>
                            <a class="doc-link" href="/case/B3?format=pdf" target="_blank">📋 Case PDF</a>
                            <a class="doc-link" href="/scoring" target="_blank">⚖️ Judge Scoring</a>
                            <a class="doc-link" href="/static/decks/B3_Slide_Deck.pptx" target="_blank">📊 Slide Deck</a>
                        </div>
                    </div>
                    <div class="case-card">
                        <div class="case-id">B4</div>
                        <div class="doc-links">
                            <a class="doc-link" href="/case/B4" target="_blank">📄 Case Page</a>
                            <a class="doc-link" href="/case/B4?format=pdf" target="_blank">📋 Case PDF</a>
                            <a class="doc-link" href="/scoring" target="_blank">⚖️ Judge Scoring</a>
                            <a class="doc-link" href="/static/decks/B4_Slide_Deck.pptx" target="_blank">📊 Slide Deck</a>
                        </div>
                    </div>
                </div>
            </div>

            <!-- PALS -->
            <div class="round-group round-pals">
                <div class="round-label"><div class="dot dot-pals"></div> PALS — Prelims Round 2</div>
                <div class="case-grid">
                    <div class="case-card">
                        <div class="case-id">P1</div>
                        <div class="doc-links">
                            <a class="doc-link" href="/case/P1" target="_blank">📄 Case Page</a>
                            <a class="doc-link" href="/case/P1?format=pdf" target="_blank">📋 Case PDF</a>
                            <a class="doc-link" href="/scoring" target="_blank">⚖️ Judge Scoring</a>
                            <a class="doc-link" href="/static/decks/P1_Slide_Deck.pptx" target="_blank">📊 Slide Deck</a>
                        </div>
                    </div>
                    <div class="case-card">
                        <div class="case-id">P2</div>
                        <div class="doc-links">
                            <a class="doc-link" href="/case/P2" target="_blank">📄 Case Page</a>
                            <a class="doc-link" href="/case/P2?format=pdf" target="_blank">📋 Case PDF</a>
                            <a class="doc-link" href="/scoring" target="_blank">⚖️ Judge Scoring</a>
                            <a class="doc-link" href="/static/decks/P2_Slide_Deck.pptx" target="_blank">📊 Slide Deck</a>
                        </div>
                    </div>
                    <div class="case-card">
                        <div class="case-id">P3</div>
                        <div class="doc-links">
                            <a class="doc-link" href="/case/P3" target="_blank">📄 Case Page</a>
                            <a class="doc-link" href="/case/P3?format=pdf" target="_blank">📋 Case PDF</a>
                            <a class="doc-link" href="/scoring" target="_blank">⚖️ Judge Scoring</a>
                            <a class="doc-link" href="/library" target="_blank">🎭 Actor Script</a>
                            <a class="doc-link" href="/static/decks/P3_Slide_Deck.pptx" target="_blank">📊 Slide Deck</a>
                        </div>
                    </div>
                    <div class="case-card">
                        <div class="case-id">P4</div>
                        <div class="doc-links">
                            <a class="doc-link" href="/case/P4" target="_blank">📄 Case Page</a>
                            <a class="doc-link" href="/case/P4?format=pdf" target="_blank">📋 Case PDF</a>
                            <a class="doc-link" href="/scoring" target="_blank">⚖️ Judge Scoring</a>
                            <a class="doc-link" href="/static/decks/P4_Slide_Deck.pptx" target="_blank">📊 Slide Deck</a>
                        </div>
                    </div>
                </div>
            </div>

            <!-- SEMI-FINALS -->
            <div class="round-group round-sf">
                <div class="round-label"><div class="dot dot-sf"></div> Semi-Finals — Round 1 &amp; Round 2</div>
                <div class="case-grid">
                    <div class="case-card">
                        <div class="case-id">R1a</div>
                        <div class="doc-links">
                            <a class="doc-link" href="/library" target="_blank">📄 Case Page</a>
                            <a class="doc-link" href="/library" target="_blank">📋 Case PDF</a>
                            <a class="doc-link" href="/library" target="_blank">⚖️ Judge Scoring</a>
                            <a class="doc-link" href="/library" target="_blank">📊 Slide Deck</a>
                        </div>
                    </div>
                    <div class="case-card">
                        <div class="case-id">R1b</div>
                        <div class="doc-links">
                            <a class="doc-link" href="/library" target="_blank">📄 Case Page</a>
                            <a class="doc-link" href="/library" target="_blank">📋 Case PDF</a>
                            <a class="doc-link" href="/library" target="_blank">⚖️ Judge Scoring</a>
                            <a class="doc-link" href="/library" target="_blank">📊 Slide Deck</a>
                        </div>
                    </div>
                    <div class="case-card">
                        <div class="case-id">R2a</div>
                        <div class="doc-links">
                            <a class="doc-link" href="/library" target="_blank">📄 Case Page</a>
                            <a class="doc-link" href="/library" target="_blank">📋 Case PDF</a>
                            <a class="doc-link" href="/library" target="_blank">⚖️ Judge Scoring</a>
                            <a class="doc-link" href="/library" target="_blank">📊 Slide Deck</a>
                        </div>
                    </div>
                    <div class="case-card">
                        <div class="case-id">R2b</div>
                        <div class="doc-links">
                            <a class="doc-link" href="/library" target="_blank">📄 Case Page</a>
                            <a class="doc-link" href="/library" target="_blank">📋 Case PDF</a>
                            <a class="doc-link" href="/library" target="_blank">⚖️ Judge Scoring</a>
                            <a class="doc-link" href="/library" target="_blank">📊 Slide Deck</a>
                        </div>
                    </div>
                </div>
            </div>

            <!-- FINALS -->
            <div class="round-group round-fn">
                <div class="round-label"><div class="dot dot-fn"></div> Finals</div>
                <div class="case-grid case-grid-2">
                    <div class="case-card">
                        <div class="case-id">F1 — PICU: Crashing on the Vent</div>
                        <div class="doc-links">
                            <a class="doc-link" href="/case/F1" target="_blank">📄 Case Page</a>
                            <a class="doc-link" href="/case/F1?format=pdf" target="_blank">📋 Case PDF</a>
                            <a class="doc-link" href="/scoring" target="_blank">⚖️ Judge Scoring</a>
                            <a class="doc-link" href="/static/decks/F1_Slide_Deck.pptx" target="_blank">📊 Slide Deck</a>
                        </div>
                    </div>
                    <div class="case-card">
                        <div class="case-id">F2 — PEM: Tense Abdomen</div>
                        <div class="doc-links">
                            <a class="doc-link" href="/case/F2" target="_blank">📄 Case Page</a>
                            <a class="doc-link" href="/case/F2?format=pdf" target="_blank">📋 Case PDF</a>
                            <a class="doc-link" href="/scoring" target="_blank">⚖️ Judge Scoring</a>
                            <a class="doc-link" href="/static/decks/F2_Slide_Deck.pptx" target="_blank">📊 Slide Deck</a>
                        </div>
                    </div>
                </div>
            </div>

        </div><!-- end cases-section -->
    </div><!-- end body -->

    <!-- ══ FOOTER ══ -->
    <div class="footer">
        <i class="fas fa-heart"></i> PediSTARS · 11th Annual National Conference on Healthcare Simulation · Bengaluru 2026
        <span class="sep">|</span> Supported by DHEE Hospitals &amp; VASA
        <br style="display:block;margin-top:6px">
        <span style="opacity:0.5">SIMWARS 2026 · Scoring Engine v8.0</span>
    </div>

</div>
</body>
</html>
"""

@app.route('/home')
def portal():
    return PORTAL_TEMPLATE
@app.route('/simwars-2026-participant-info.html')
def participant_info():
        return send_file(os.path.join(os.path.dirname(__file__), 'simwars-2026-participant-info.html'))

@app.route('/draw-sheet')
def draw_sheet():
    locked = require_role()
    if locked: return locked
    # Judges see the official draw only on event-day morning (22 Aug, 06:00 IST).
    # Organisers always have access.
    if session.get('role') == 'judge' and _ist_now() < datetime(2026, 8, 22, 6, 0):
        return ("<div style='font-family:sans-serif;max-width:420px;margin:20vh auto;text-align:center;"
                "color:#2a0a44;'><h2>\U0001F3B2 Official Draw Sheet</h2>"
                "<p>The draw is revealed to judges on the morning of the event "
                "&mdash; <b>22 Aug 2026, 06:00 IST</b>. Come back then!</p></div>"), 403
    return DRAW_SHEET_TEMPLATE

DRAW_SHEET_TEMPLATE = r"""<!DOCTYPE html>
<html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>SIM WARS 2026 — Official Draw Sheet</title>
<style>
*{margin:0;padding:0;box-sizing:border-box;}
body{font-family:Inter,-apple-system,'Segoe UI',sans-serif;background:#f4f1f7;color:#2a0a44;padding:1.5rem;}
.sheet{max-width:900px;margin:0 auto;background:#fff;border-radius:14px;padding:1.6rem 1.8rem;box-shadow:0 8px 30px rgba(42,10,68,.12);}
h1{font-size:1.35rem;font-weight:900;text-align:center;}
.sub{text-align:center;color:#6a5a72;font-size:.85rem;margin:.3rem 0 1.2rem;}
.grid{display:grid;grid-template-columns:1fr 1fr;gap:.9rem;}
.g{border-radius:12px;padding:.9rem 1rem;}
.g h2{font-size:.75rem;font-weight:900;letter-spacing:.12em;margin-bottom:.5rem;}
.t{display:flex;justify-content:space-between;gap:.6rem;font-size:.86rem;padding:.35rem 0;border-bottom:1px dashed rgba(0,0,0,.08);}
.t:last-child{border-bottom:none;}
.t .n{font-weight:900;}
.t .tbd{color:#9aa0ae;font-style:italic;}
.sched{margin-top:1.3rem;}
.sched h2{font-size:1rem;font-weight:900;margin-bottom:.5rem;}
table{width:100%;border-collapse:collapse;font-size:.82rem;}
th,td{border:1px solid #e8e0ee;padding:.45rem .6rem;text-align:left;}
th{background:#f4effa;font-size:.72rem;letter-spacing:.08em;text-transform:uppercase;}
.badge{display:inline-block;font-size:.68rem;font-weight:900;padding:.15rem .55rem;border-radius:40px;}
.toolbar{display:flex;justify-content:space-between;align-items:center;max-width:900px;margin:0 auto 1rem;}
.btn{background:#2a0a44;color:#fff;font-weight:800;font-size:.82rem;padding:.55rem 1.1rem;border-radius:60px;text-decoration:none;border:none;cursor:pointer;}
@page{size:A4;margin:10mm;}
@media print{body{background:#fff;padding:0;}.toolbar{display:none;}.sheet{box-shadow:none;padding:0;}}
</style></head><body>
<div class="toolbar"><a class="btn" href="/register">&#8592; Home</a><button class="btn" onclick="window.print()">🖨️ Print / Save PDF</button></div>
<div class="sheet">
<h1>🎲 SIM WARS 2026 — Official Draw &amp; Colour Groups</h1>
<div class="sub">Drawn 21 Aug 2026 · 16 teams · Prelims Fri 22 Aug, DHEE Hospitals · <span id="stamp"></span></div>
<div class="grid" id="groups">Loading teams…</div>
<div class="sched">
<h2>Preliminary Round Schedule</h2>
<table><thead><tr><th>Round</th><th>Time</th><th>Room A · PALS</th><th>Room B · BLS</th></tr></thead>
<tbody>
<tr><td><b>R1</b></td><td>10:00 – 11:30</td><td><span class="badge" style="background:#fdecea;color:#b02a37;">RED</span> draws 1–4</td><td><span class="badge" style="background:#fef9c3;color:#a16207;">YELLOW</span> draws 13–16</td></tr>
<tr><td><b>R2</b></td><td>11:45 – 13:15</td><td><span class="badge" style="background:#e8f7f1;color:#0d8a72;">GREEN</span> draws 5–8</td><td><span class="badge" style="background:#fdecea;color:#b02a37;">RED</span> draws 1–4</td></tr>
<tr><td><b>R3</b></td><td>14:00 – 15:30</td><td><span class="badge" style="background:#eef6fc;color:#2f6fd1;">BLUE</span> draws 9–12</td><td><span class="badge" style="background:#e8f7f1;color:#0d8a72;">GREEN</span> draws 5–8</td></tr>
<tr><td><b>R4</b></td><td>15:30 – 17:00</td><td><span class="badge" style="background:#fef9c3;color:#a16207;">YELLOW</span> draws 13–16</td><td><span class="badge" style="background:#eef6fc;color:#2f6fd1;">BLUE</span> draws 9–12</td></tr>
</tbody></table>
</div>
</div>
<script>
var INTAKE='https://script.google.com/macros/s/AKfycbwyhhMFVFHo1IQCQw97Hzfv8PSWgXWgChtmbvqnXDtTA7pKExHsvBM3hL5rL279EjOM/exec';
var GROUPS=[{name:'RED',draws:[1,2,3,4],bg:'#fdecea',fg:'#b02a37'},{name:'GREEN',draws:[5,6,7,8],bg:'#e8f7f1',fg:'#0d8a72'},{name:'BLUE',draws:[9,10,11,12],bg:'#eef6fc',fg:'#2f6fd1'},{name:'YELLOW',draws:[13,14,15,16],bg:'#fef9c3',fg:'#a16207'}];
function esc(s){return String(s==null?'':s).replace(/[&<>"']/g,function(c){return{'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'}[c];});}
document.getElementById('stamp').textContent='Generated '+new Date().toLocaleString('en-IN',{timeZone:'Asia/Kolkata'})+' IST';
fetch(INTAKE).then(function(r){return r.json();}).then(function(d){
 var byDraw={};(d.teams||[]).forEach(function(t){byDraw[parseInt(t.drawNumber,10)]=t;});
 var h='';
 GROUPS.forEach(function(g){
  h+='<div class="g" style="background:'+g.bg+';"><h2 style="color:'+g.fg+';">'+g.name+' GROUP</h2>';
  g.draws.forEach(function(n){
   var t=byDraw[n];
   h+='<div class="t"><span><span class="n" style="color:'+g.fg+';">'+n+'</span> · '+(t?esc(t.teamName):'<span class="tbd">TBD</span>')+'</span>'
     +(t&&t.members?'<span style="font-size:.72rem;color:#6a5a72;max-width:55%;text-align:right;">'+esc(Array.isArray(t.members)?t.members.join(', '):t.members)+'</span>':'')+'</div>';
  });
  h+='</div>';
 });
 document.getElementById('groups').innerHTML=h;
}).catch(function(){document.getElementById('groups').innerHTML='<div style="color:#b02a37;">Could not load team intake — check connection.</div>';});
</script>
</body></html>"""

@app.route('/simwars-2026-questionnaire.html')
def questionnaire():
    import re as _re
    with open(os.path.join(os.path.dirname(__file__), 'simwars-2026-questionnaire.html'), encoding='utf-8') as _f:
        _html = _f.read()
    # Gate is open to ALL teams 1-16 regardless of which HTML copy is deployed
    _html = _re.sub(r'const allowedTeams = \[[^\]]*\];',
                    'const allowedTeams = [1,2,3,4,5,6,7,8,9,10,11,12,13,14,15,16];', _html)
    _resp = app.response_class(_html, mimetype='text/html')
    _resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    return _resp

@app.route('/stay')
@app.route('/simwars-2026-stay-brochure.html')
def stay_brochure():
    return send_file(os.path.join(os.path.dirname(__file__), 'simwars-2026-stay-brochure.html'))

@app.route('/whats-new')
@app.route('/simwars-2026-whats-new.html')
def whats_new():
    return send_file(os.path.join(os.path.dirname(__file__), 'simwars-2026-whats-new.html'))

@app.route('/judges-guide')
@app.route('/simwars-2026-judges-bengaluru.html')
def judges_guide():
    return send_file(os.path.join(os.path.dirname(__file__), 'simwars-2026-judges-bengaluru.html'))

@app.route('/mentors')
@app.route('/simwars-2026-mentor-flyer.html')
def mentor_flyer():
    return send_file(os.path.join(os.path.dirname(__file__), 'simwars-2026-mentor-flyer.html'))

@app.route('/watch')
@app.route('/simwars-2026-spectator-flyer.html')
def spectator_flyer():
    return send_file(os.path.join(os.path.dirname(__file__), 'simwars-2026-spectator-flyer.html'))

def _ensure_questionnaire_table(conn):
    conn.execute('CREATE TABLE IF NOT EXISTS questionnaire (id INTEGER PRIMARY KEY AUTOINCREMENT, team TEXT, payload TEXT, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)')

@app.route('/api/questionnaire', methods=['POST'])
def submit_questionnaire():
    data = request.get_json(silent=True) or {}
    if not data:
        return jsonify({'ok': False, 'error': 'empty submission'}), 400
    with get_db() as conn:
        _ensure_questionnaire_table(conn)
        conn.execute('INSERT INTO questionnaire (team, payload) VALUES (?, ?)',
                     (str(data.get('teamNumber', 'unknown')), _json.dumps(data)))
        conn.commit()
    return jsonify({'ok': True})

@app.route('/questionnaire-responses')
def questionnaire_responses():
    locked = require_organiser()
    if locked: return locked
    with get_db() as conn:
        _ensure_questionnaire_table(conn)
        rows = conn.execute('SELECT team, payload, created_at FROM questionnaire ORDER BY created_at').fetchall()
    if request.args.get('format') == 'json':
        return jsonify([{'team': r['team'], 'created_at': str(r['created_at']), 'payload': _json.loads(r['payload'])} for r in rows])
    items = ''.join(
        '<tr><td>%s</td><td>%s</td><td><details><summary>view</summary><pre>%s</pre></details></td></tr>'
        % (r['team'], r['created_at'], _json.dumps(_json.loads(r['payload']), indent=1).replace('<', '&lt;'))
        for r in rows)
    return ('<!DOCTYPE html><html><head><meta charset="utf-8"><title>Questionnaire Responses</title>'
            '<style>body{font-family:-apple-system,Segoe UI,Roboto,sans-serif;padding:24px;color:#1b2430;}'
            'table{border-collapse:collapse;width:100%%;}td,th{border:1px solid #e3e7ee;padding:8px 12px;font-size:14px;text-align:left;vertical-align:top;}'
            'pre{white-space:pre-wrap;font-size:12px;max-height:300px;overflow:auto;background:#f7f8fb;padding:8px;border-radius:6px;}'
            'h1{font-size:22px;}a{color:#d81b7a;font-weight:700;}</style></head><body>'
            '<h1>Pre-Event Questionnaire — %d response(s)</h1>'
            '<p><a href="/questionnaire-responses?format=json">Download all as JSON</a></p>'
            '<table><tr><th>Team</th><th>Submitted</th><th>Responses</th></tr>%s</table></body></html>'
            % (len(rows), items))


if __name__ == '__main__':
    import os as _os
    port = int(_os.environ.get('PORT', 8080))
    app.run(host='0.0.0.0', debug=False, port=port)
